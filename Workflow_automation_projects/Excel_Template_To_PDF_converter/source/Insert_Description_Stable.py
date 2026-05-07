import os
# keep your original QT plugin path line if you need it for your environment
os.environ["QT_QPA_PLATFORM_PLUGIN_PATH"] = r"C:\Users\Jayesh_Bhat\miniconda3\Library\plugins\platforms"

from docx import Document as dok
from copy import deepcopy
from pathlib import Path
from typing import List
import shutil
import subprocess
import sys
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def open_with_default_app(path: Path) -> None:
    """Open file with the OS default application (cross-platform)."""
    p = str(path)
    if sys.platform.startswith("win"):
        import os as _os

        _os.startfile(p)
    elif sys.platform == "darwin":
        subprocess.run(["open", p], check=False)
    else:
        subprocess.run(["xdg-open", p], check=False)


def prompt_yes_no(prompt: str) -> bool:
    """Ask user y/n until valid response returned."""
    while True:
        ans = input(prompt + " [y/n]: ").strip().lower()
        if ans in ("y", "yes"):
            return True
        if ans in ("n", "no"):
            return False
        print("Please answer 'y' or 'n'.")


def subsection_with_explanation(obj, Heading, Description):
    """
    Creates a subsection with a heading and paragraph from description
    """
    obj.write("\\section*{" + escape_special_chars(Heading) + "}\n")
    obj.write(escape_special_chars(Description) + "\n")
    
    
def table_description(obj, table_list, table_para_list, table_col_para_list):
    """
    Creates a description object using the elements read from the table.
    Writes a clean \item[...] for the first line of each cell and
    writes the rest (same-cell remainder or the paired next cell) as content.
    - If the paired next cell contains '=' we create an itemize from those lines.
    - If the paired next cell contains 'For this study' we indent it with \hspace{1cm}.
    """
    obj.write("\\begin{description}\n")

    # step through pairs (k, k+1)
    for k in range(0, max(0, len(table_list) - 1), 2):
        cell = table_list[k] or ""
        paired = table_list[k + 1] or ""

        # heading is the first line of the current cell
        parts = cell.split("\n", 1)
        heading = parts[0].strip()
        remainder = parts[1].strip() if len(parts) > 1 and parts[1].strip() else None

        # write the item label using only the heading
        obj.write("\\item[\\small " + escape_special_chars(heading) + "]\\mbox{}\n")

        # if current cell had a remainder (same cell description), write it
        if remainder:
            # add an explicit linebreak in LaTeX if we will also write paired content
            if paired.strip() and "=" in paired:
                obj.write(escape_special_chars(remainder) + "\n")
            else:
                obj.write(escape_special_chars(remainder) + "\\\\"+"\n")

        # handle the paired (next) cell
        if paired.strip():
            paired_text = paired.strip()

            # if the paired cell contains equals-sign list items, itemize them
            if "=" in paired_text:
                # split into lines, strip empty lines, and itemize only those lines
                items = [line.strip() for line in paired_text.splitlines() if line.strip()]
                new_list = itemize_function(items)
                for para_item in new_list:
                    obj.write(para_item)
            # if it mentions "For this study" indent it
            elif all(good_token in paired_text for good_token in ["For", "this", "study"]) and all(bad_token not in paired_text for bad_token in ["=", ":"]):
                obj.write("\\hspace{1 cm} " + escape_special_chars(paired_text) + "\n")
            else:
                # fallback: write as plain text
                obj.write(escape_special_chars(paired_text) + "\n")

    obj.write("\\end{description}\n")


def itemize_function(table_paragraph_list):
    """
    A function that lists iterative table elements using the itemize function of latex
    """
    # keep debug prints if you find them useful; can be turned to logger.debug()
    logger.debug("itemize input: %s", table_paragraph_list)
    itemize_list = []
    itemize_list.append("\\begin{itemize}\n")
    # mutate a local list copy to avoid surprising caller sides
    working = list(table_paragraph_list)
    while len(working) > 0:
        if "=" in working[0]:
            item_str = "\t" + "\\item " + escape_special_chars(working[0]) + "\n"
            itemize_list.append(item_str)
            working.pop(0)
            logger.debug('"=" read; remaining: %s', working)
            if (len(working) <= 0) or (":" in working[0]):
                logger.debug("itemize loop broken; remaining: %s", working)
                break
        else:
            working.pop(0)

    itemize_list.append("\\end{itemize}" + "\\" + "\n")
    return itemize_list


def table_heading(col_1_para_list, num):
    """Return heading from first column paragraphs list by index."""
    if num < 0 or num >= len(col_1_para_list):
        return ""
    return col_1_para_list[num]


def table_subdescription(col_1_para_list, num):
    """Return subdescription (next item) from first column paragraphs list by index."""
    if (num + 1) < 0 or (num + 1) >= len(col_1_para_list):
        return ""
    return col_1_para_list[num + 1]


def escape_special_chars(textWUnd: str) -> str:
    """
    A function that handles latex special characters
    """
    if textWUnd is None:
        return ""
    new_str = ""
    for ch in textWUnd:
        if ch == "_":
            new_str += "\\_ "
        elif ch == "&":
            new_str += "\\& "
        elif ch == "$":
            new_str += "\\$ "
        elif ch == "%":
            new_str += "\\% "
        elif ch == "{":
            new_str += "\\{ "
        elif ch == "}":
            new_str += "\\} "
        elif ch == "#":
            new_str += "\\# "
        elif ch == "~":
            new_str += "\\textasciitilde "
        elif ch == "^":
            new_str += "\\textasciicircum "
        elif ch == "\\":
            new_str += "\\textbackslash "
        else:
            new_str += ch
    return new_str


def clean_list(table_list: List[str]) -> List[str]:
    """
    Remove empty or whitespace-only strings from list
    """
    return [elem for elem in table_list if elem and elem.strip()]


# ------------------------ Helper steps extracted from the big function ------------------------


def create_description_docx_from_template(template_path: Path, dest_folder: Path, tex_source: Path) -> Path:
    """Copy template docx to description folder and return new docx path."""
    dest_folder.mkdir(parents=True, exist_ok=True)
    new_name = tex_source.stem + "_description.docx"
    dest_path = dest_folder / new_name
    # prefer filesystem copy to deepcopy of loaded Document (safer, simpler)
    shutil.copyfile(template_path, dest_path)
    logger.info("Copied template to %s", dest_path)
    return dest_path


def parse_docx_document(docx_path: Path):
    """Parse docx and return paragraphs list and tables as flattened cell texts."""
    document = dok(str(docx_path))
    all_paragraphs = document.paragraphs
    logger.info("No of paragraphs read: %d", len(all_paragraphs))

    names_explanation_list = []
    for para in all_paragraphs:
        if para.text and para.text.strip():
            names_explanation_list.append(para.text.strip())

    all_tables = document.tables
    logger.info("No of tables read: %d", len(all_tables))

    dict_table = {}
    for obj_num, obj in enumerate(all_tables, start=1):
        list_for_table = f"table_{obj_num}_list"
        dict_table[list_for_table] = []
        if obj_num == len(all_tables):
            dict_table[f"table_{obj_num}_para_list"] = []
            dict_table[f"table_{obj_num}_col_1_para_list"] = []

    # Accessing text from each cell in each table
    for table_num, table in enumerate(all_tables, start=1):
        key = f"table_{table_num}_list"
        for row in table.rows:
            for cell in row.cells:
                dict_table[key].append(cell.text)

    # For last table also collect per-paragraph lists
    for table_num, table in enumerate(all_tables, start=1):
        if table_num == len(all_tables):
            para_key = f"table_{table_num}_para_list"
            col1_key = f"table_{table_num}_col_1_para_list"
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        dict_table[para_key].append(para.text)
                        if ":" not in para.text or "=" not in para.text:
                            dict_table[col1_key].append(para.text)

    return names_explanation_list, dict_table


def write_description_tex_from_parsed(
    names_explanation_list: List[str], dict_table: dict, description_tex_path: Path, original_num_tables: int
):
    """Write the description .tex file using the parsing results."""
    # prepare keys
    key_list = list(dict_table.keys())
    list_of_tables = key_list[0 : (len(key_list) - 2)] if len(key_list) >= 2 else []
    table_X_para_list = key_list[-2] if len(key_list) >= 2 else None
    table_X_col_1_para_list = key_list[-1] if len(key_list) >= 1 else None

    clean_table_X_col_1_para_list = []
    if table_X_col_1_para_list in dict_table:
        clean_table_X_col_1_para_list = clean_list(dict_table[table_X_col_1_para_list])

    # remove title if present
    if names_explanation_list:
        title = names_explanation_list.pop(0)
    else:
        title = ""

    # table explanation items (last N*2 items — matches original logic)
    if original_num_tables > 0:
        table_explanation_list = names_explanation_list[(len(names_explanation_list) - (original_num_tables * 2)) :]
    else:
        table_explanation_list = []

    with description_tex_path.open("w", encoding="utf8") as TxF:
        if title:
            TxF.write("\\section*{\\huge " + escape_special_chars(title) + "}" + "\n")

        # normal sections
        z = 0
        while z < len(names_explanation_list):
            if names_explanation_list[z] in table_explanation_list:
                break
            # guard for out-of-range
            heading = names_explanation_list[z] if z < len(names_explanation_list) else ""
            body = names_explanation_list[z + 1] if (z + 1) < len(names_explanation_list) else ""
            TxF.write("\\section*{" + escape_special_chars(heading) + "}\n")
            TxF.write(escape_special_chars(body) + "\n")
            z += 2

        # now add the table explanation sections
        i = 0
        j = 0
        while i < len(table_explanation_list) and j < len(list_of_tables):
            subsection_with_explanation(TxF, table_explanation_list[i], table_explanation_list[i + 1])
            table_description(TxF, dict_table[list_of_tables[j]], dict_table[table_X_para_list], clean_table_X_col_1_para_list)
            j += 1
            i += 2


def insert_input_into_tex(original_tex: Path, description_tex: Path, final_tex: Path):
    """Insert \\input{description_tex} after \\begin{document} and write final file."""
    with original_tex.open("r", encoding="utf8") as ftex:
        lines = ftex.readlines()

    inserted = False
    new_lines = []
    desc_path_unix = str(description_tex).replace("\\", "/")
    input_line = r"\input{" + f"{desc_path_unix}" + "}\n" + r"\newpage" + "\n"

    for line in lines:
        new_lines.append(line)
        if not inserted and r"\begin{document}" in line:
            new_lines.append(input_line)
            inserted = True

    if not inserted:
        raise RuntimeError(r"\begin{document} not found in the source .tex file")

    with final_tex.open("w", encoding="utf8") as f_new:
        f_new.writelines(new_lines)


# ------------------------ Main orchestrator (refactored) ------------------------


def insert_description_file(tex_file_without_description: str, template_docx: str = None, description_folder: str = None) -> str:
    """
    High-level orchestrator. Returns path to final tex file (string).
    Parameters:
      - tex_file_without_description: path to the .tex that currently lacks description.
      - template_docx: optional path to standard template (if None, uses your previous default).
      - description_folder: optional folder for description files (if None, uses your previous default).
    """
    tex_src = Path(tex_file_without_description)

    # default paths preserved from your original file if not provided
    if template_docx is None:
        template_docx = r"D:\Code\Software_test_sample_data\Dev_Proj_5__ExcelToPdfConverter\DESCPR_files\std_word_template\Standard_Latex_Description_File.docx"
    if description_folder is None:
        description_folder = r"D:\Code\Software_test_sample_data\Dev_Proj_5__ExcelToPdfConverter\DESCPR_files"

    template_path = Path(template_docx)
    desc_folder = Path(description_folder)

    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    # 1) create description docx copy
    new_docx = create_description_docx_from_template(template_path, desc_folder, tex_src)

    # 2) open file for user to edit and wait for confirmation
    open_with_default_app(new_docx)
    if not prompt_yes_no("Have you setup the word description file?"):
        raise RuntimeError("User exited before completing the description docx")

    # 3) parse docx into paragraphs and table dict
    names_explanation_list, dict_table = parse_docx_document(new_docx)
    original_num_tables = len(list(k for k in dict_table.keys() if k.startswith("table_") and k.endswith("_list")))

    # 4) write description tex file
    doc_tex_filename = desc_folder / (tex_src.stem + "_description.tex")
    write_description_tex_from_parsed(names_explanation_list, dict_table, doc_tex_filename, original_num_tables)

    # 5) insert into original tex to create final tex
    final_tex = tex_src.with_name(tex_src.stem + "_final.tex")
    insert_input_into_tex(tex_src, doc_tex_filename, final_tex)

    logger.info("Created final tex: %s", final_tex)
    return str(final_tex)


# Keep the old name for backward compatibility
def Insert_description_file(tex_file_without_description: str):
    return insert_description_file(tex_file_without_description)
