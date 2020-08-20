from docx import Document as dok

def main():
    # Read text from each paragraphs
    document = dok('Documentation_JBTest_4_EDITED.docx')

    all_paragraphs = document.paragraphs
    print("\n\nNo of paragraphs read: " + str(len(all_paragraphs)))


    # List for collecting all paragraphs - Section headings and corressponding descriptions are picked from this list
    names_explanation_list = []

    for num, para in enumerate(all_paragraphs, start=1):
        if para.text == '':
            continue
        print(str(num) + '-' + para.text)
        names_explanation_list.append(para.text)


    # Read text from each cell in each Table
    all_tables = document.tables
    print("\n\nNo of tables read: " + str(len(all_tables)))


    table_1_list = []
    table_2_list = []
    table_3_list = []
    table_4_list = []
    table_4_para_list = []


    for table_num, table in enumerate(all_tables):
        if table_num == 0:
            for row in table.rows:
                for cell in row.cells:
                    table_1_list.append(cell.text)
        elif table_num == 1:
            for row in table.rows:
                for cell in row.cells:
                    table_2_list.append(cell.text)
        elif table_num == 2:
            for row in table.rows:
                for cell in row.cells:
                    table_3_list.append(cell.text)
        elif table_num == 3:
            for row in table.rows:
                for cell in row.cells:
                    table_4_list.append(cell.text)
                    for para in cell.paragraphs:
                            table_4_para_list.append(para.text)


    print(names_explanation_list)
    print(len(names_explanation_list))

    print(f"""
            table_1_list contains:
            {table_1_list}

            table_2_list contains: 
            {table_2_list}

            table_3_list contains
            {table_3_list}
            
            table_4_heading_list contains
            {table_4_list}

            table_4_para_list contains
            {table_4_para_list}

            """)


    list_of_tables = [table_1_list, table_2_list, table_3_list, table_4_list]
    table_lenghts_list = [len(table_1_list), len(table_2_list), len(table_3_list), len(table_4_list)]
    print(table_lenghts_list)


    title = names_explanation_list.pop(0)
    section_1_heading = names_explanation_list.pop(0)
    section_1_description = names_explanation_list.pop(0)
    section_2_heading = names_explanation_list.pop(0)
    section_2_description = names_explanation_list.pop(0)


    with open('Doc_04.tex', 'w') as TxF:
        TxF.write('\\documentclass[a4paper, 12pt]{article}\n')
        TxF.write('\\begin{document}\n')
        TxF.write('\\section{' + title + '}' + '\\' + '\n')
        TxF.write('\\subsection{' + section_1_heading + '}\n')
        TxF.write(section_1_description + '\\' + '\\' + '\n')
        TxF.write('\\subsection{' + section_2_heading + '}\n')
        TxF.write(section_2_description + '\\' + '\\' + '\n')

        i = 0
        j = 0
        while i < len(names_explanation_list) and j < len(list_of_tables):
            print('i starts with:' + str(i))
            print('j starts with:' + str(j))
            subsection_with_explanation(TxF, names_explanation_list[i], names_explanation_list[i+1])
            table_description(TxF, list_of_tables[j], table_4_para_list)
            j += 1
            i += 2
            print('j ends with:' + str(j))
            print('i ends with:' + str(i))
        TxF.write('\\end{document}\n')


def subsection_with_explanation(obj, Heading, Description):
    obj.write('\subsubsection{' + Heading + '}\n')
    obj.write(Description + '\\' + '\n')


def table_description(obj, table_list, table_para_list):
    k = 0
    obj.write('\\begin{description}\n')
    while k < len(table_list):
        if table_list[len(table_list)-1] and '=' in table_list[k + 1] and ':' in table_list[k]:
            obj.write('\\item[' + escape_special_chars(table_list[k]) + ']\mbox{}\n')
            new_list = itemize_function(table_para_list)
            for para_item in new_list:
                obj.write(para_item)
        else:
            obj.write('\\item[' + escape_special_chars(table_list[k]) + ']' + escape_special_chars(table_list[k + 1]) + '\n')
        k += 2
    obj.write('\\end{description}' + '\\' + '\n')


def itemize_function(table_paragraph_list):
    print(table_paragraph_list)
    itemize_list = []
    itemize_list.append('\\begin{itemize}\n')
    table_paragraph_list.pop(0)
    while len(table_paragraph_list) > 1:
        if '=' in table_paragraph_list[0]:
            item_str = ('\t' + '\\item ' + table_paragraph_list[0] + '\n')
            itemize_list.append(item_str)
            table_paragraph_list.pop(0)
            print('"=" read: ')
            print(table_paragraph_list)
            if ':' in table_paragraph_list[0]:
                table_paragraph_list.pop(0)
                print('loop broken!')
                print(table_paragraph_list)
                break
        elif '' in table_paragraph_list[1]:
            table_paragraph_list.pop(0)
            print('" '' " removed!')
            print(table_paragraph_list)
            print(itemize_list)
            table_paragraph_list.pop(0)
            continue
    itemize_list.append('\\end{itemize}' + '\\' + '\n')
    return itemize_list


def escape_special_chars(textWUnd):
    if '_' in textWUnd:
        txt_split_1 = textWUnd.split('_')
        txt_join_1 = '\_'.join(txt_split_1)
        return txt_join_1
    elif '°' in textWUnd:
        txt_split_2 = textWUnd.split('°')
        txt_join_2 = '\°'.join(txt_split_2)
        return txt_join_2
    return textWUnd


if __name__ == '__main__':
    main()
