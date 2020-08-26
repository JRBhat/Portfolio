from docx import Document as dok


def main():
    # specify the file to parse
    document = dok("19.0510-23_image_overview.docx")

    # collecting all paragraph objects and assigning it to a variable
    all_paragraphs = document.paragraphs
    print("\n\nNo of paragraphs read: " + str(len(all_paragraphs)))

    # Empty List for collecting each paragraph text generated from paragraphs objects
    names_explanation_list = []

    for num, para in enumerate(all_paragraphs, start=1):
        if para.text == "":
            continue
        print(str(num) + "-" + para.text)
        names_explanation_list.append(para.text)

    # Reads text from each cell in each Table
    all_tables = document.tables

    # Create a dictionary which creates empty lists with indexing according to number of table objects stored in 'all_tables' variable
    dict_table = {}

    for obj_num, obj in enumerate(all_tables, start=1):
        list_for_table = "table_" + str(obj_num) + "_list"

        dict_table[list_for_table] = []
        if obj_num == len(all_tables):
            list_for_para = "table_" + str(obj_num) + "_para_list"
            list_for_col = "table_" + str(obj_num) + "_col_1_para_list"
            dict_table[list_for_para] = []
            dict_table[list_for_col] = []

    print()
    print(dict_table)
    print("\n\nNo of tables read: " + str(len(all_tables)))

    # Accessing text from each cell in each table
    for table_num, table in enumerate(all_tables, start=1):
        for row in table.rows:
            for cell in row.cells:
                dict_table["table_" + str(table_num) + "_list"].append(cell.text)

    for table_num, table in enumerate(all_tables, start=1):
        if table_num == len(all_tables):
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        dict_table["table_" + str(table_num) + "_para_list"].append(
                            para.text
                        )
                        if ":" in para.text or "=" not in para.text:
                            dict_table[
                                "table_" + str(table_num) + "_col_1_para_list"
                            ].append(para.text)

    ### Printing for viewing list elements
    print(names_explanation_list)
    print(len(names_explanation_list))

    for key, value in dict_table.items():
        print("\n\n" + "List name:" + key, end="\n")
        print(value)

    key_list = list(dict_table.keys())
    print()
    print("This is the key list")
    print(key_list)

    list_of_tables = key_list[0 : (len(key_list) - 2)]
    table_X_para_list = key_list[-2]
    table_X_col_1_para_list = key_list[-1]

    ### Printing lengths of each list
    """ 
        print(list_of_tables)
        print()
        print(table_X_para_list)
        print()
        print(table_X_col_1_para_list) """

    # Removes spaces and null strings from the list
    clean_table_X_col_1_para_list = clean_list(dict_table[table_X_col_1_para_list])
    print()
    print(clean_table_X_col_1_para_list)

    # First few elements of the list occur every time, so they are standardized and removed each time
    title = names_explanation_list.pop(0)
    section_1_heading = names_explanation_list.pop(0)
    section_1_description = names_explanation_list.pop(0)
    section_2_heading = names_explanation_list.pop(0)
    section_2_description = names_explanation_list.pop(0)

    # Creating a new tex document and writing the created standardized objects into it
    with open("TEST_19.0510-23_image_overview.tex", "w") as TxF:
        TxF.write("\\documentclass[a4paper, 12pt]{article}\n")
        TxF.write("\\begin{document}\n")
        TxF.write("\\section{" + title + "}" + "\\" + "\n")
        TxF.write("\\subsection{" + section_1_heading + "}\n")
        TxF.write(section_1_description + "\\" + "\\" + "\n")
        TxF.write("\\subsection{" + section_2_heading + "}\n")
        TxF.write(section_2_description + "\\" + "\\" + "\n")

        i = 0
        j = 0
        while i < len(names_explanation_list) and j < len(list_of_tables):
            print("i starts with:" + str(i))
            print("j starts with:" + str(j))
            subsection_with_explanation(
                TxF, names_explanation_list[i], names_explanation_list[i + 1]
            )
            table_description(
                TxF,
                dict_table[list_of_tables[j]],
                dict_table[table_X_para_list],
                clean_table_X_col_1_para_list,
            )
            j += 1
            i += 2
            print("j ends with:" + str(j))
            print("i ends with:" + str(i))
        TxF.write("\\end{document}\n")


def subsection_with_explanation(obj, Heading, Description):
    """ Creates a subsection with a heading and paragraph from description """
    obj.write("\subsubsection{" + Heading + "}\n")
    obj.write(Description + "\\" + "\n")


def table_description(obj, table_list, table_para_list, table_col_para_list):
    """Creates a Latex description object using the elements read from the table. 
    It also creates an itemize Latex object for listing out table sub-elements """

    k = 0
    obj.write("\\begin{description}\n")
    while k < len(table_list):
        if (
            table_list[len(table_list) - 1]
            and "=" in table_list[k + 1]
            and ":" in table_list[k]
        ):
            obj.write(
                "\\item["
                + escape_special_chars(table_heading(table_col_para_list, k))
                + "]\mbox{}\n"
            )
            obj.write(table_subdescription(table_col_para_list, k) + "\\" + "\\" + "\n")
            new_list = itemize_function(table_para_list)
            for para_item in new_list:
                obj.write(para_item)
        else:
            obj.write(
                "\\item["
                + escape_special_chars(table_list[k])
                + "]"
                + escape_special_chars(table_list[k + 1])
                + "\n"
            )
        k += 2
    obj.write("\\end{description}" + "\\" + "\n")


def itemize_function(table_paragraph_list):
    """A function that lists iterating table elements using the itemize latex object"""
    print(table_paragraph_list)
    itemize_list = []
    itemize_list.append("\\begin{itemize}\n")
    while len(table_paragraph_list) > 0:
        if "=" in table_paragraph_list[0]:
            item_str = "\t" + "\\item " + table_paragraph_list[0] + "\n"
            itemize_list.append(item_str)
            table_paragraph_list.pop(0)
            print('"=" read: ')
            print(table_paragraph_list)
            if (len(table_paragraph_list) <= 0) or (":" in table_paragraph_list[0]):
                print("loop broken!")
                print(table_paragraph_list)
                break
        else:
            table_paragraph_list.pop(0)

    itemize_list.append("\\end{itemize}" + "\\" + "\n")
    return itemize_list


def table_heading(col_1_para_list, num):
    """Takes in the heading as a paragraph from the first column of the table and prits it out as Subsectin title"""
    return col_1_para_list[num]


def table_subdescription(col_1_para_list, num):
    """Takes in the description from the first column of the table and prits it out as Subsectin description"""
    return col_1_para_list[num + 1]


def escape_special_chars(textWUnd):
    """A function that adds a backshlash infront of special characters such as '_' and '°' to escape them"""
    if "_" in textWUnd:
        txt_split_1 = textWUnd.split("_")
        txt_join_1 = "\_".join(txt_split_1)
        return txt_join_1
    elif "°" in textWUnd:
        txt_split_2 = textWUnd.split("°")
        txt_join_2 = "\°".join(txt_split_2)
        return txt_join_2
    return textWUnd


def clean_list(table_list):
    """A function that takes in a list and removes spaces and null strings from the list"""
    cl_list = []
    for elem in table_list:
        if elem != "" and elem != " ":
            cl_list.append(elem)
    return cl_list


if __name__ == "__main__":
    main()
