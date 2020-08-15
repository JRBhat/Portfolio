from docx import Document as dok
from pylatex import Document, Section,  Description, NewLine
from pylatex.utils import bold


def main():
    # Read text from each paragraphs
    document = dok('Documentation_JBTest_2.docx')
    all_paragraphs = document.paragraphs
    print("\n\nNo of paragraphs read: " + str(len(all_paragraphs)))

    # List that collects all paragraphs - Used for section names and explanation arguments later
    names_explanation_list = []

    for num, para in enumerate(all_paragraphs, start=1):
        if para.text == '':
            continue
        print(str(num) + '-' + para.text)
        names_explanation_list.append(para.text)

    # Read text from each cell in each Table
    all_tables = document.tables
    print("\n\nNo of tables read: " + str(len(all_tables)))


    table_1_list = []
    table_2_list = []
    table_3_list = []

    for table_num, table in enumerate(all_tables):
        if table_num == 0:a
            for row in table.rows:
                for cell in row.cells:
                    table_1_list.append(cell.text)
        elif table_num == 1:
            for row in table.rows:
                for cell in row.cells:
                    table_2_list.append(cell.text)
        elif table_num == 2:
            for row in table.rows:
                for cell in row.cells:
                    table_3_list.append(cell.text)



    print(names_explanation_list)
    print(len(names_explanation_list))

    print(f"""
            table_1_list contains:
            {table_1_list}

            table_2_list contains: 
            {table_2_list}

            table_3_list contains
            {table_3_list}
            """)

    list_of_tables = [table_1_list,'', table_2_list,'', table_3_list, '']
    table_lenghts_list = [len(table_1_list), len(table_2_list), len(table_3_list)]
    print(table_lenghts_list)
    doc = Document()

    names_explanation_list.remove(names_explanation_list[0])
    print(len(names_explanation_list))
    print(names_explanation_list)

    i = 0
    j = 0

    while i < len(names_explanation_list):
        print('i starts with:' + str(i))
        create_sections(doc, names_explanation_list[i], names_explanation_list[i+1])
        while j < len(list_of_tables[i]):
            print('j starts with:' + str(j))
            create_description(doc, list_of_tables[i][j], (list_of_tables[i][j+1]))
            j += 2
            print('j ends with:' + str(j))
        i += 2
        j = 0
        print('i ends with:' + str(i))

    doc.generate_pdf('not_fun')
    doc.generate_tex('not_fun')

def create_sections(doc_obj, name, expl):
    with doc_obj.create(Section(bold(name))):
        doc_obj.append(expl)

def create_description(doc_obj, first, last):
    with doc_obj.create(Description()) as desc:
        desc.add_item(first, last)

    """
    for i in range(0,len(names_explanation_list), 2):
        with doc.create(Section(bold(names_explanation_list[i + 1]))):
            doc.append(names_explanation_list[i + 2])
    """

    """    # List that contains text from each table cell
        table_cell_list = []
        for table_num, table in enumerate(all_tables):
            print('All cells in table' + '-' + str(table_num) + ' are as follows: \n')

            for row_num, row in enumerate(table.rows, start=1):
                for cell_num, cell in enumerate(row.cells):
                    print('  Paragraphs[row-' + str(row_num) + ',cell-' + str(cell_num) + ']: ')
                    for para_num, para in enumerate(cell.paragraphs):
                        if para.text == '':
                            continue
                        print('  ' + str(para_num) + '-' + para.text + '\n')
                        table_cell_list.append(para.text)
            print('\n------------------------------------------------------')"""

if __name__ == '__main__':
    main()

