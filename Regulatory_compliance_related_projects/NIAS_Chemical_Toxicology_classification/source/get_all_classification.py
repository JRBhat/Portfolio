
from docx import Document
import re 
import openpyxl
import os
from utilityFuncs.util import leading_cas_number
import logging

# Configure logger
"""
# Example usage
logger = logging.getLogger(__name__)  # Get logger instance

logger.debug("This is a debug message")  # Low-level info for debugging
logger.info("This is an info message")  # General info
logger.warning("This is a warning message")  # Warning
logger.error("This is an error message")  # Error
logger.critical("This is a critical message")  # Critical error
"""
logging.basicConfig(
    level=logging.DEBUG,  # Set log level (DEBUG, INFO, WARNING, ERROR, CRITICAL)
    format="%(asctime)s - %(levelname)s - %(message)s",  # Log format
    handlers=[
        logging.FileHandler("app.log"),  # Save logs to a file
        logging.StreamHandler()  # Print logs to console
    ]
)


def get_template_type_and_language(paragraphs):# Template material? data/reports/Test/excluded/template--MATERIAL
    
    template_type = ""
    # Determine the Language and get the template type  from table 1b footnote- Food/Kosmetik
    if 'Ergebnisse' in paragraphs[0].text:
        language = 'DE'
        for p in paragraphs:
            if "mg/kg " in p.text and "Lebensmittel" in p.text:
                    template_type = "L"
                    break
            elif "mg/kg " in p.text and "Kosmetikum" in p.text:
                    template_type = "K"
                    break
                
    elif 'Results' in paragraphs[0].text:
        language = 'EN'
        for p in paragraphs:
            if "mg/kg" in p.text and "food" in p.text:
                template_type = "L"
                break
            elif "mg/kg" in p.text and "cosmetic" in p.text:
                template_type = "K"
                break
     
    else:
        language = 'Unknown'
        template_type = 'Unknown'
    
    return language, template_type


def parse_classifications_from_doc(reports_path, docx_file):
        
    # Extracting doc file contents
    doc = Document(f'{reports_path}/{docx_file}')
    tables = doc.tables
    paragraphs = doc.paragraphs
    
    language, template_type = get_template_type_and_language(paragraphs)
    print(f'{docx_file} - {language} - {template_type}')
    
    
    clf_txt = ""
    classf_dict = {}
    for tno, t in enumerate(tables):
        if tno > 4:
            break
        print(f"Table {tno + 1}")
        
        
        for row in range(1, len(t.rows)): #skip first row headers
            for col in range(0, len(t.columns)): 
                  
                # classification data
                
                if (tno==1 or tno==3) and col == 0:
                    
                    clf_txt, clf_cas, clf_substance_name = get_classification(t, row)

                    classf_dict[(clf_cas.lstrip('0'), clf_substance_name)] = {"Substance": clf_substance_name, 
                                                                              "CL": clf_txt, 
                                                                              "language": language, 
                                                                              "template_type": template_type
                                                                              }
    return classf_dict



        
def get_classification(t, row):
    
    clf_txt = ""
    clf_cas  =""
    clf_substance_name = ""
    
    clf_data = [(t.cell(0, colm._index).text, colm._index) 
                for colm in t.columns
                if "Classification" in t.cell(0, colm._index).text 
                or "Einordnung" in t.cell(0, colm._index).text
                ]
    if len(clf_data) > 0:
        
        clf_indx = clf_data[0][1]
        clf_substance_name = t.cell(row, clf_indx-2).text
        clf_cas = t.cell(row, clf_indx-1).text
        clf_txt = t.cell(row, clf_indx).text
        
    return clf_txt, clf_cas, clf_substance_name



def save_dict_to_excel(data_dict, output_file):
    
    column_names_specified = ["CAS", 
                              "Substance_DE",
                              "Substance_EN", 
                              "CL_DE", 
                              "CL_EN", 
                              "language", 
                              "template_type"]
                                  
    # Check if the file exists
    if os.path.exists(output_file):
        wb = openpyxl.load_workbook(output_file)
        # Use the "Doc_data" sheet if it exists, otherwise create it and add a header row
        if "Doc_data" in wb.sheetnames:
            ws = wb["Doc_data"]
        else:
            ws = wb.create_sheet("Doc_data")
            ws.append(column_names_specified)
    else:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Doc_data"
        ws.append(column_names_specified)

    # Append each dictionary's data as a new row.
    # To ensure consistent column order, we explicitly extract values in the header order.
    for key, data in data_dict.items():
        if data["language"] == "DE":
            
            row = [key[0], 
                   data["Substance"],
                   "", 
                   data["CL"], 
                   "", 
                   data["language"],
                   data["template_type"]]
            ws.append(row)
        else:
            row = [key[0],
                   "", 
                   data["Substance"],  
                   "", 
                   data["CL"],
                   data["language"],
                   data["template_type"]]
            ws.append(row)
    wb.save(output_file)
    print(f"Excel file saved as {output_file}")
    
def main(TESTING = False):
        
    # Example usage
    # report_name = r"SAMPLE_REPORT_I01.docx"
    # path = os.environ.get("REPORTS_DIR", "data/reports/test_classification")
    path = os.environ.get("REPORTS_DIR", "data/reports")
    #
    for report_name in os.listdir(path):
        if report_name.endswith("docx"):
            # data = parse_footnotes_from_doc(path, report_name)
            data = parse_classifications_from_doc(path, report_name)
            if TESTING:
                out_path = os.path.join(path, rf"{report_name}_output2_clf4.xlsx")   
                save_dict_to_excel(data, out_path)
            else:
                save_dict_to_excel(data, os.path.join("\\".join(path.split("\\")[:-3]), "output_ALL_v8_20250223_clf4.xlsx"))
    
if __name__ == "__main__":
    main(TESTING=False)