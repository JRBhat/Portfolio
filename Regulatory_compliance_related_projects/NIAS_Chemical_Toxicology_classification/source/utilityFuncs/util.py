"""
utilityFuncs/util.py
====================
Cross-cutting helper functions used throughout the NIAS reporting pipeline.

Categories
----------
Word / table helpers
    ``set_table_cell_text``, ``set_table_cell_text_fts``,
    ``get_text_and_superscript``, ``get_cl_and_inci``
    Provide a thin wrapper over the ``python-docx`` cell API to enforce
    consistent font (Arial 8 pt) and to handle superscript footnote markers.

Chemical identifier helpers
    ``validate_cas`` – check-digit validation for CAS registry numbers.
    ``format_cas`` – left-pad the leading numeric segment to 6 digits.
    ``capitalize_first_alpha`` – capitalize first letter of a chemical name.
    ``get_iupac_name`` – return the language-appropriate IUPAC name from a
    Substance object, falling back to the trivial name.
    ``leading_cas_number`` – regex match callback that zero-pads CAS prefixes.

Database mapping loaders
    ``load_cramer_TTC_mapping`` – reads the ``LEGEND-Cramer`` sheet and returns
    a dict of Cramer class → {Long_Name, TTC} pairs.
    ``load_fcm_group_mapping`` – reads the ``CL_Groups`` sheet and returns a
    dict of group key → bilingual metadata.

External API helpers
    ``query_pubchem`` – three-step PubChem REST query:
    (1) resolve CID, (2) fetch SMILES/IUPAC/formula, (3) extract CAS from
    synonyms.
    ``query_chemspider_api`` – RSC ChemSpider API helper (requires API key in
    ``RSC_CHEMSPIDER_API_KEY`` environment variable; used for common-name resolution).

Logging
-------
A dedicated ``general_errors`` logger writes ERROR-level messages to
``errors.log`` and to the console.  Debug-level URL messages are also emitted
(visible when the log level is lowered).
"""

import re
import pandas as pd
import numpy as np
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
import logging
from dotenv import load_dotenv
import os
import pandas as pd
from typing import Dict, Any

import requests
import re
import logging


# ----------------------------
# Logger Setup
# ----------------------------
formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')

general_logger = logging.getLogger('general_errors')
general_logger.setLevel(logging.ERROR)
gen_file_handler = logging.FileHandler('errors.log', mode='w')
gen_file_handler.setFormatter(formatter)
general_logger.addHandler(gen_file_handler)
gen_console_handler = logging.StreamHandler()
gen_console_handler.setFormatter(formatter)
general_logger.addHandler(gen_console_handler)


def format_cas(cas, format_str='06d'):
    try:
        leading_group_old = re.match(r'^(\d+)', cas).group(1)
        leading_group_new = format(int(leading_group_old), format_str)
        return cas.replace(leading_group_old, leading_group_new)
    except AttributeError:
        input("Wait")


def set_table_cell_text(table, row, col, text, center=False):
    """
    Sets text in a table cell and optionally aligns it to the center.

    Args:
        table (docx.Table): The table object.
        row (int): The row index of the cell.
        col (int): The column index of the cell.
        text (str): Text to set in the cell.
        center (bool): Whether to center-align the text. Default is False.
    """

    if type(text) is list:
        inci_indices = [i for i in range(len(text)) if 'INCI' in text[i]]
        inci_text = None
        if len(inci_indices) > 0:
            inci_text = text[inci_indices[0]]
        text = [x for x in text if not 'INCI' in x]
        if inci_text is not None:
            text.append(inci_text)
        text = '\n'.join(text)

    table.cell(row, col).text = text
    paragraph = table.cell(row, col).paragraphs[0]
    paragraph.paragraph_format.space_after = 0
    if center:
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs
    font = run[0].font
    font.name = 'Arial'
    font.size = Pt(8)


def set_table_cell_text_fts(table, row, col, text, fts):
    """
    Sets text and footnotes in a table cell. Footnotes are displayed as superscripts.

    Args:
        table (docx.Table): The table object.
        row (int): The row index of the cell.
        col (int): The column index of the cell.
        text (str): Text to display in the cell.
        fts (dict): Dictionary of footnotes for the cell.
    """
    paragraph = table.cell(row, col).paragraphs[0]
    paragraph.add_run()
    paragraph.add_run()


    runs = paragraph.runs

    font = runs[0].font
    font.name = 'Arial'
    font.size = Pt(8)
    runs[0].text = text

    font = runs[1].font
    font.name = 'Arial'
    font.size = Pt(8)
    font.superscript = True
    for ft in sorted(list(fts.keys())):
        runs[1].text += ft
        runs[1].text += ', '
    runs[1].text = runs[1].text.rstrip(', ')

    table.cell(row, col).paragraph = paragraph




def get_iupac_name(language, s):
    if language == 'DE':
        if s.iupac_name_de is not None:
            iupac_name = s.iupac_name_de
        else:
            iupac_name = s.name
    elif language == 'EN':
        if s.iupac_name_en is not None:
            iupac_name = s.iupac_name_en
        else:
            iupac_name = s.name
    return iupac_name


# war 07d
def leading_cas_number(match):
    number = int(match.group(1))
    return format(number, '06d') #convert into 6 digits by adding leading zeros




def load_cramer_TTC_mapping(excel_path: str, sheet_name: str = "LEGEND-Cramer",
                            decimal_sep: str = ",") -> Dict[str, Dict[str, Any]]:
    """
    Reads the given Excel file and sheet, expecting columns:
      - 'Cramer' (string key)
      - 'TTC'    (float, with comma decimal by default)
      - 'LongName' (string)

    Returns a dict of the form:
      {
        'I':       {'Long_Name': 'Substances of Cramer class I',           'TTC': 0.03},
        'II':      {'Long_Name': 'Substances of Cramer class II',          'TTC': 0.009},
        'III':     {'Long_Name': 'Substances of Cramer class III',         'TTC': 0.0015},
        'OP':      {'Long_Name': 'Organophosphates and carbamates',        'TTC': 0.0003},
        'GENOTOX': {'Long_Name': 'Substances with structural alerts of genotoxicity', 'TTC': 2.5e-06},
      }
    """
    # Read the sheet, parsing comma or specified decimal separator
    df = pd.read_excel(
        excel_path,
        sheet_name=sheet_name,
        dtype={'Cramer': str, 'TTC': float, 'LongName': str},
        decimal=decimal_sep
    )

    # Build and return the mapping
    return {
        row.Cramer: {
            'Long_Name': row.LongName,
            'TTC':       row.TTC
        }
        for row in df.itertuples(index=False)
    }



def load_fcm_group_mapping(excel_path: str, sheet_name: str = "CL_Groups") -> Dict[int, Dict[str, str]]:
    """
    Reads the given Excel file and sheet, expecting columns:
      - 'Group_Key' (integer key)
      - 'Name_DE' (string)
      - 'Name_EN' (string)
      - 'CL_DE' (string)
      - 'CL_EN' (string)

    Returns a dict of the form:
      {
        9: {'Names_DE': '...', 'Name_EN': '...', 'CL_DE': '...', 'CL_EN': '...'},
        13: {'Names_DE': '...', 'Name_EN': '...', 'CL_DE': '...', 'CL_EN': '...'},
        ...
      }
    """
    # Read the sheet
    df = pd.read_excel(
        excel_path,
        sheet_name=sheet_name,
        dtype={
            'Group_Key': str,
            'Name_DE': str,
            'Name_EN': str,
            'CL_DE': str,
            'CL_EN': str,
            'Type_Flag': str
        }
    )

    # Build and return the mapping
    return {
        row.Group_Key: {
            'Name_DE': row.Name_DE,
            'Name_EN': row.Name_EN,
            'CL_DE': row.CL_DE,
            'CL_EN': row.CL_EN,
            'Type_Flag': row.Type_Flag
        }
        for row in df.itertuples(index=False)
    }



def get_text_and_superscript(cell):
    """
    Extracts text and superscript annotations from a table cell.

    Args:
        cell (docx.Table.Cell): The table cell to process.

    Returns:
        tuple: A tuple containing the main text and superscript text.
    """
    name = ''
    ft = ''
    for run in cell.paragraphs[0].runs:
        if not run.font.superscript:
            name += run.text
        else:
            if not any([x.strip().isdigit() for x in run.text.split(',')]):
                ft += run.text
            else:
                pass  # Add to name?
    return (name, ft)



def get_cl_and_inci(text):
    """
    Extracts classification (CL) and INCI information from a text block.

    Args:
        text (str): The input text to parse.

    Returns:
        tuple: A tuple containing a list of classifications and INCI text.
    """
    lines = text.split('\n')
    cls = []
    inci = ''
    for line in lines:
        if 'INCI Name' in line:
            #inci = line.replace('INCI Name:', '').strip()
            cls.append(line.strip())
        else:
            cls.append(line.strip())
    return (cls, inci)



def merge_info(substances, new_substances):
    """
    Merges new substance information into an existing dictionary of substances.

    Args:
        substances (dict): Existing dictionary of substances.
        new_substances (dict): New substances to merge into the dictionary.
    """
    for s in new_substances.values():
        if s.cas in substances.keys():
            s_db = substances[s.cas]
            if len(s.cl) > 0:
                for cl in s.cl:
                    if cl not in substances[s.cas].cl:
                        substances[s.cas].cl.append(cl)
            if len(s.ft) > 0:
                for ft in s.ft:
                    if ft not in substances[s.cas].ft:
                        substances[s.cas].ft.append(ft)
            if s.inci != '':
                substances[s.cas].inci = s.inci
        else:
            substances[s.cas] = s


def capitalize_first_alpha(s):
    """Capitalize the first alphabetic character in the string."""
    return next((s[:i] + c.upper() + s[i+1:] for i, c in enumerate(s) if c.isalpha()), s)

def validate_cas(cas):
    """Validate a CAS number using its check digit."""
    cas_pattern = re.compile(r'^\d{2,7}-\d{2}-\d$')
    if cas_pattern.match(cas):
        positions = np.arange(9, -1, -1)
        ziffern = np.array(list(f"{int(''.join(cas.split('-'))):010d}"), dtype=int)
        betrag = positions[:-1] * ziffern[:-1]
        return int(np.sum(betrag) % 10) == ziffern[-1]
    return False


def query_pubchem(identifier, query_type="cas", text_changed=None):
    """Query PubChem using an identifier (CAS, IUPAC name, or SMILES)."""
    identifier = identifier.strip()

    if query_type.lower() in ["cas", "iupac", "name"]:
        input_type = "name"
    elif query_type.lower() == "smiles":
        input_type = "smiles"
    else:
        raise ValueError("Invalid query_type. Use 'cas', 'iupac', or 'smiles'.")

    # Step 1: Resolve CID
    cid_url = f"https://pubchem.ncbi.nlm.nih.gov/rest/pug/compound/{input_type}/{identifier}/cids/JSON"
    try:
        general_logger.debug(f"Fetching CID from URL: {cid_url}")
        cid_response = requests.get(cid_url, verify=True)
        cid_response.raise_for_status()
        cid_data = cid_response.json()
        cids = cid_data.get("IdentifierList", {}).get("CID", [])
        if not cids:
            if text_changed:
                text_changed.emit(f"⚠️ PubChem: No CID found for identifier {identifier}")
            general_logger.error(f"PubChem: No CID found for identifier: {identifier}")
            return None
        cid = cids[0]

        # Step 2: Fetch properties
        prop_url = (
            f"https://pubchem.ncbi.nlm.nih.gov/rest/pug/compound/cid/{cid}/"
            "property/ConnectivitySMILES,IUPACName,MolecularFormula/JSON"
        )
        general_logger.debug(f"Fetching properties from URL: {prop_url}")
        prop_response = requests.get(prop_url, verify=True)
        prop_response.raise_for_status()
        prop_data = prop_response.json()
        properties = prop_data.get("PropertyTable", {}).get("Properties", [])
        if not properties:
            if text_changed:
                text_changed.emit(f"⚠️ PubChem: No properties found for CID {cid}")
            general_logger.error(f"PubChem: No properties found for CID: {cid}")
            return None

        result = {
            "Identifier": identifier,
            "SMILES": properties[0].get("ConnectivitySMILES"),
            "IUPAC": properties[0].get("IUPACName"),
            "MolecularFormula": properties[0].get("MolecularFormula"),
            "CAS": None
        }

        # Step 3: Fetch synonyms to find CAS
        syn_url = f"https://pubchem.ncbi.nlm.nih.gov/rest/pug/compound/cid/{cid}/synonyms/JSON"
        general_logger.debug(f"Fetching synonyms from URL: {syn_url}")
        syn_response = requests.get(syn_url, verify=True)
        syn_response.raise_for_status()
        syn_data = syn_response.json()

        synonyms = syn_data.get("InformationList", {}).get("Information", [{}])[0].get("Synonym", [])
        cas_pattern = re.compile(r'^\d{2,7}-\d{2}-\d$')
        for syn in synonyms:
            if cas_pattern.match(syn):
                result["CAS"] = syn
                break

        return result

    except requests.RequestException as e:
        if text_changed:
            text_changed.emit(f"⚠️ PubChem query failed for identifier {identifier}: {e}")
        general_logger.error(f"PubChem query failed for identifier {identifier}: {e}")
    except Exception as e:
        if text_changed:
            text_changed.emit(f"⚠️ Unexpected error for identifier {identifier}: {e}")
        general_logger.exception(f"Unexpected error for identifier {identifier}: {e}")
    return None


# Load ChemSpider API key from environment variable.
# Set RSC_CHEMSPIDER_API_KEY in your environment or a .env file before running.
load_dotenv()  # Loads from a .env file in the current working directory if present
API_KEY = os.getenv("RSC_CHEMSPIDER_API_KEY")
BASE_URL = "https://api.rsc.org/compounds/v1"


def query_chemspider_api(identifier: str) -> dict | None:
    """
    Query the RSC Compounds v1 API for CAS, IUPAC, SMILES, and English & German names.
    Returns a dict or None on error / not found.

    Requires the RSC_CHEMSPIDER_API_KEY environment variable to be set.
    """
    identifier = identifier.strip()
    # 1) Pick filter endpoint

    endpoint = "filter/name"
    payload  = {"name": identifier}
    headers = {
        "apikey": API_KEY,
        "Content-Type": "application/json"
    }

    try:
        # 2) POST to filter → get queryId
        r = requests.post(f"{BASE_URL}/{endpoint}", json=payload, headers=headers)
        r.raise_for_status()
        qid = r.json().get("queryId")
        if not qid:
            return None

        # 3) GET results → list of CSIDs
        r = requests.get(f"{BASE_URL}/filter/{qid}/results", headers=headers)
        r.raise_for_status()
        res = r.json().get("results", [])
        if res:
            csid = res[0]
            if not csid:
                return None
            else:
                # 4) Fetch details (which now includes structures, CAS, plus often synonyms)
                r = requests.get(f"{BASE_URL}/records/{csid}/details?fields=CommonName", headers=headers)
                r.raise_for_status()
                details = r.json()
                common_name  = details.get("commonName")

                return {
                    "Identifier": identifier,
                    "CSID":       csid,
                    "CommonName": common_name,
                }

    except requests.RequestException as exc:
        print(f"ChemSpider query failed: {exc}")
        return None


def delete_substance(signal, s):
    while True:
        if signal == "y":
            del s
            print(f"deleted {s.name}")
            break
        if signal == "n":
            print(f"{s.name} NOT deleted")
            break
