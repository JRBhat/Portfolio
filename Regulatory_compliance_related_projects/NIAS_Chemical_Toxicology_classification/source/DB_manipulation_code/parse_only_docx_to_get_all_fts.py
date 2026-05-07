
from docx import Document
import re 
import openpyxl
import os
from utilityFuncs.util import leading_cas_number
import logging

# Configure logger
"""
# Example usage
logger = logging.getLogger(__name__)  # Get logger instance

logger.debug("This is a debug message")  # Low-level info for debugging
logger.info("This is an info message")  # General info
logger.warning("This is a warning message")  # Warning
logger.error("This is an error message")  # Error
logger.critical("This is a critical message")  # Critical error
"""
logging.basicConfig(
    level=logging.DEBUG,  # Set log level (DEBUG, INFO, WARNING, ERROR, CRITICAL)
    format="%(asctime)s - %(levelname)s - %(message)s",  # Log format
    handlers=[
        logging.FileHandler("app.log"),  # Save logs to a file
        logging.StreamHandler()  # Print logs to console
    ]
)


def parse_footnotes_from_doc(reports_path, docx_file):
        
    # Extracting doc file contents
    doc = Document(f'{reports_path}/{docx_file}')
    tables = doc.tables
    paragraphs = doc.paragraphs
    

    language, template_type = get_template_type_and_language(paragraphs)
    print(f'{docx_file} - {language}')


    doc_ft = get_footnote_text_from_paragraphs(paragraphs)
    
    logging.info(len(doc_ft))
    logging.info(docx_file)

    
    doc_substances = {}
    Footnote_dict = {}

    
    substance_name = ""
    cramer = ""
    cas = ""
    clf_txt = ""
    ft_indices = ""
    table_nr_order = []
    
    for tno, t in enumerate(tables):
        if tno > 4:
            break
        print(f"Table {tno + 1}")
        

        
        for row in range(1, len(t.rows)): #skip first row headers
            for col in range(0, len(t.columns)): 

                
                # superscript, substance name and cas
                if ((tno==0 or tno==2) and 
                    col == 0 and # first column 
                    len(t.cell(row, col).paragraphs)<=2 and # one paragraph - means single line - (skips multiple lines such as those for Retention times one under other)
                    len(t.cell(row, col).paragraphs[0].runs)>1 and  # multiple runs - becasue footnotes have a different format/style causing a different run
                    True in [r.font.superscript for r in t.cell(row, col).paragraphs[0].runs] # one of the runs must be a superscript
                ):

                    # substance name
                    substance_name = get_substance_name_from_paragraph(t, row, col)
                    print(f"Table {tno} -  Substance name: {substance_name}")
                    
                    # footnote indices
                    ft_indices = "".join([r.text for r in t.cell(row, col).paragraphs[0].runs if r.font.superscript])
                    print(f"Table {tno} -  FT_superscript: {ft_indices}")
                    
                    # cas
                    cas = get_cas(t, row, col,substance_name)
                    print(f"Table {tno} - CAS: {cas}")
                    
                    # cramer
                    cramer = get_cramer_class(t, row)
                    print(f"Table {tno} - Cramer: {cramer}")
                    table_nr_order.append(tno)
                    table_no = tno
                    Footnote_dict[cas] = [substance_name, ft_indices, cas, cramer, table_no, '']
                  
                # classification data
                
                if (tno==1 or tno==3) and len(Footnote_dict.keys()) > 0 and col == 0:
                    
                    clf_txt, clf_cas, clf_substance_name = get_classification(t, row)
                    for cas, val in Footnote_dict.items():
                        
                        if (clf_cas.strip() == cas.strip() 
                            and val[-1] == "" and 
                            (strip_name(clf_substance_name) == strip_name(val[0]))
                        ):
                            Footnote_dict[cas][-1] = clf_txt
                            break
                        
                        if ("$" in cas and 
                            val[-1] == "" and
                            (strip_name(clf_substance_name) == strip_name(cas[1:])) and 
                            (strip_name(clf_substance_name) == strip_name(val[0]))
                        ):
                            
                            Footnote_dict[cas][-1] = clf_txt
                            if clf_cas != "-" and re.search(r"[0-9]*-[0-9]{2}-[0-9]{1}", clf_cas)!= None:
                                Footnote_dict[cas][2] = clf_cas
                            else:
                                Footnote_dict[cas][2] = "$" + strip_name(clf_substance_name)
                            break
                        
    Footnote_dict = validate_dictkey_cas(Footnote_dict)                    
    doc_substances = map_tableno_to_substance_groups(table_nr_order, doc_substances, Footnote_dict)
    
    report_path = f'{reports_path}/{docx_file}'
    return join_data(doc_ft, doc_substances, report_path, language, template_type)

def validate_dictkey_cas(Footnote_dict):
    new_dict = {}
    
    for k, v in Footnote_dict.items():
        if (k != v[2] and re.search(r"[0-9]*-[0-9]{2}-[0-9]{1}", v[2])!= None):
            # change the key to the cas saved in the values, only if a valid cas is detected
            new_dict[v[2]] = v
        else:
            new_dict[k] = v
    return new_dict

def strip_name(name_str):
    return name_str.replace(",", "").replace("-", "").replace("[", "").replace("]", "").replace(".", "").replace(" ", "").replace(")", "").replace("(", "").lower()

def compare_stripped_names(str1, str2):
    clean_str1 = strip_name(str1)
    clean_str2 = strip_name(str2)
    print(clean_str1, clean_str2, clean_str1 == clean_str2)
    return clean_str1 == clean_str2

def map_tableno_to_substance_groups(table_nr_order, doc_substances, Footnote_dict):
    
    table_nr_order = set(table_nr_order)
    for nr in table_nr_order:
        doc_substances[nr + 1] = {k:v for k, v in Footnote_dict.items() if v[-2] == nr}
        
    return doc_substances

def get_footnote_text_from_paragraphs(paragraphs):
    
    # Extract Footnotes
    doc_ft = {} 
    p_table = 0
    for _, p in enumerate(paragraphs):
        if 'Tab. 1' in p.text:
            p_table = 1
            if p_table not in doc_ft.keys():
                doc_ft[p_table] = {}
        elif 'Tab. 2' in p.text:
            p_table = 2
            if p_table not in doc_ft.keys():
                doc_ft[p_table] = {}
        elif 'Tab. 3' in p.text:
            p_table = 3
            if p_table not in doc_ft.keys():
                doc_ft[p_table] = {}
        elif 'Tab. 4' in p.text:
            p_table = 4
            if p_table not in doc_ft.keys():
                doc_ft[p_table] = {}
        # elif 'Tab. A' in p.text:
        #     p_table = None

        
        is_ft = False
        ft_index = ''
        ft_text = ''
        for _, r in enumerate(p.runs): #In Microsoft Word (when using Python with python-docx), 
            # paragraph.runs refers to a list of Run objects within a Paragraph. 
            # Each Run represents a contiguous block of text with the same formatting (e.g., bold, italics, font size, etc.).
            if r.font.superscript is not None:
                if ft_index == '':
                    ft_index = r.text.strip()
                    if ft_index != '':
                        is_ft = True
            else:
                if is_ft:
                    ft_text += r.text

        if is_ft and p_table is not None:
            print(f'{p_table}.{ft_index} - {ft_text.strip()}')
            try:
                doc_ft[p_table][ft_index] = ft_text.strip() 
            except KeyError:
                continue
                
    return doc_ft

def get_template_type_and_language(paragraphs):# Template material? data/reports/Test/excluded/template--MATERIAL
    
    template_type = ""
    # Determine the Language and get the template type  from table 1b footnote- Food/Kosmetik
    if 'Ergebnisse' in paragraphs[0].text:
        language = 'DE'
        for p in paragraphs:
            if "mg/kg " in p.text and "Lebensmittel" in p.text:
                    template_type = "L"
                    break
            elif "mg/kg " in p.text and "Kosmetikum" in p.text:
                    template_type = "K"
                    break
                
    elif 'Results' in paragraphs[0].text:
        language = 'EN'
        for p in paragraphs:
            if "mg/kg" in p.text and "food" in p.text:
                template_type = "L"
                break
            elif "mg/kg" in p.text and "cosmetic" in p.text:
                template_type = "K"
                break
     
    else:
        language = 'Unknown'
        template_type = 'Unknown'
    
    return language, template_type
        
def get_classification(t, row):
    
    clf_txt = ""
    clf_cas  =""
    clf_substance_name = ""
    
    clf_data = [(t.cell(0, colm._index).text, colm._index) 
                for colm in t.columns
                if "Classification" in t.cell(0, colm._index).text 
                or "Einordnung" in t.cell(0, colm._index).text
                ]
    if len(clf_data) > 0:
        
        clf_indx = clf_data[0][1]
        clf_substance_name = t.cell(row, clf_indx-2).text
        clf_cas = t.cell(row, clf_indx-1).text
        clf_txt = t.cell(row, clf_indx).text
        
    return clf_txt, clf_cas, clf_substance_name
        
def get_substance_name_from_paragraph(t, row, col):
    # extracting substance name from paragraph
    runs_per_para = [(r.text, r.style.name, r.font.name, r.font.superscript)  for r in t.cell(row, col).paragraphs[0].runs]
    
    substance_name = "".join([l[0] for l in runs_per_para if l[-1] != True]) 
    
    if "CAS" in substance_name:
        cas = re.search(r"[0-9]*-[0-9]{2}-[0-9]{1}", substance_name).group(0)
        substance_name = substance_name.replace(f"(CAS {cas})", "")
    print(f"substance: {substance_name}")
    return substance_name

def get_cramer_class(t, row):
    
    # cramer
    cramer = ""
    cramer_data = [(t.cell(0, colm._index).text, colm._index) for colm in t.columns if "Cramer" in t.cell(0, colm._index).text]
    if len(cramer_data) > 0:
        col_indx = cramer_data[0][1]
        col_data = t.cell(row, col_indx).text
        if  col_data in ["I", "II", "III", "-"]:
            cramer = col_data
        else :
            cramer = "n.a."
    
    return cramer

def get_cas(t, row, col, substance_name):
    cas = ""
    if "CAS" in t.cell(0, 1).text:
        try:
            cas =  re.search(r"[0-9]*-[0-9]{2}-[0-9]{1}", t.cell(row, col+1).text).group(0)
        except AttributeError:
            cas = "$"+ strip_name(substance_name)
            
    if "CAS" in t.cell(row, col).text:
        cas = re.search(r"[0-9]*-[0-9]{2}-[0-9]{1}", t.cell(row, col).text).group(0)
    else:
        cas = "$"+ strip_name(substance_name)

    return cas

def join_data(dict1, dict2, report_path, language, template_type):
    # Assuming 'dict1' and 'dict2' are as provided, and we use dict1[3] for these indices.
    joined_data = {}
    for _, tbl_data in dict2.items():
        if len(tbl_data) > 0:
            for cas, (substance_name, ft_indices, cas_number, cramer, tbl_no, classification) in tbl_data.items():
                # Split the index string into individual letter codes.
                letters = [letter.strip() for letter in ft_indices.split(',')]
                # Retrieve the corresponding text from dict1[3].
                try:
                    descriptions = [dict1[tbl_no+1].get(letter, f"(No entry for '{letter}')") for letter in letters]
                    # Join the descriptions into one string.
                    full_description = str(descriptions)
                except KeyError:
                    #with open("bad_reports.log", "a+") as f:
                        # print(f"BAD_KEY found in - {report_path} for table {tbl_no} and {tbl_data}\n\n", file=f)
                    full_description = ""
                # Store or print the result.
                joined_data[cas] = {
                    "substance_name": substance_name,
                    "cas": re.sub(r'^(\d+)', leading_cas_number, cas_number),
                    "cramer": cramer,
                    "classification": classification,
                    "table_no": tbl_no+1,
                    "ft_indices": ft_indices,
                    "ft_text": full_description,
                    "language": language,
                    "template_type": template_type,
                    "path": report_path
                }
    return joined_data

def save_dict_to_excel(data_dict, output_file):
    column_names_specified = ["substance_name", "cas", "cramer", "classification", "table_no", "ft_indices", "ft_text", "language", "template_type", "path"]
    # Check if the file exists
    if os.path.exists(output_file):
        wb = openpyxl.load_workbook(output_file)
        # Use the "Doc_data" sheet if it exists, otherwise create it and add a header row
        if "Doc_data" in wb.sheetnames:
            ws = wb["Doc_data"]
        else:
            ws = wb.create_sheet("Doc_data")
            ws.append(column_names_specified)
    else:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Doc_data"
        ws.append(column_names_specified)

    # Append each dictionary's data as a new row.
    # To ensure consistent column order, we explicitly extract values in the header order.
    for _, data in data_dict.items():
            row = [
                data.get(name, "") for name in column_names_specified
            ]
            ws.append(row)
    
    wb.save(output_file)
    print(f"Excel file saved as {output_file}")
    
def main(TESTING = False):
        
    # Example usage
    # report_name = r"SAMPLE_REPORT_I01.docx"
    path = os.environ.get("REPORTS_DIR", "data/reports/test_classification")
    # path = os.environ.get("REPORTS_DIR", "data/reports")
    #
    for report_name in os.listdir(path):
        if report_name.endswith("docx"):
            data = parse_footnotes_from_doc(path, report_name)
            if TESTING:
                out_path = os.path.join(path, rf"{report_name}_output2_chk3.xlsx")   
                save_dict_to_excel(data, out_path)
            else:
                save_dict_to_excel(data, os.path.join(path, "output_ALL_v8_20250223_chk3.xlsx"))
    
if __name__ == "__main__":
    main(TESTING=False)