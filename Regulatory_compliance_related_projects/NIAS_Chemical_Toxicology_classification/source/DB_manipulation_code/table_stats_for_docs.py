import os
import pandas as pd
from docx import Document
import shutil

def parse_footnotes_from_doc(reports_path, docx_file):
    # Extracting doc file contents
    doc = Document(f'{reports_path}/{docx_file}')
    tables = doc.tables
    paragraphs = doc.paragraphs
    filename = docx_file
    path = f'{reports_path}/{docx_file}'
    
    # Extract Footnotes
    all_table_labels = []
    impt_table_labels = []
    
    for p in paragraphs:
        if 'Tab.' in p.text[:9]:
            all_table_labels.append(p.text[:9])

        if any(f'Tab. {i}' in p.text[:9] for i in range(1, 5)):
            impt_table_labels.append(p.text[:9])

    if len(all_table_labels) == 0:
        all_table_labels = ["N.A."]
        
    if len(impt_table_labels) == 0:
        impt_table_labels = ["N.A."]
            
    return filename, all_table_labels, len(tables), impt_table_labels, path

def save_data_to_excel(data, output_path):
    """
    Saves the parsed data to an Excel file.
    """
    def ensure_string(value):
        """Ensures lists are converted to a single string, keeping other values intact."""
        if isinstance(value, list):
            return ', '.join(map(str, value))  # Convert list elements to string
        return str(value)

    df = pd.DataFrame([
        {
            "filename": ensure_string(row[0]),
            "All Table Labels": ensure_string(row[1]),
            "Total Tables": ensure_string(row[2]),
            "Important Table Labels": ensure_string(row[3]),
            "path": ensure_string(row[4]),
        }
        for row in data
    ])

    with pd.ExcelWriter(output_path, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="Tables Data")


def move_excluded(all_data, excluded_path):
    
    for row in all_data:
        if 'N.A.' in row[1] or 'N.A.' in row[3]:
            print(row)
            prompt = input(f"move y/n? {row[-1]}")
            if prompt == "y":
                shutil.move(row[-1], os.path.join(excluded_path, row[0]))
        
  
def main(mode="report"):
    path = os.environ.get("REPORTS_DIR", "data/reports")
    stat_output_path = os.path.join(path, "output_tables_stats_v2.xlsx")
    excluded_path = os.path.join(path, "excluded")
    
    all_data = []
    for report_name in os.listdir(path):
        if report_name.endswith("docx"):
            data = parse_footnotes_from_doc(path, report_name)
            if [] in data or None in data or "" in data:
                input(f"{data}")
            else:
                all_data.append(data)
                print(f"{report_name} added..")

    if all_data:
        if mode == "report":
            
            save_data_to_excel(all_data, stat_output_path)  
            print(f"Saved output to {stat_output_path}")
            
        if mode == "exclude":
            move_excluded(all_data, excluded_path)
       


if __name__ == "__main__":
    main(mode="report")
    # main(mode="exclude")
