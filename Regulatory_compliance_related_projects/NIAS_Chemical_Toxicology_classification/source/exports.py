"""
exports.py
==========
Generates the final Word document (.docx) compliance report by populating a
pre-formatted DOCX template with processed substance data.

The module is the last stage of the NIAS reporting pipeline.  It receives fully
enriched substance objects (with FCM numbers, Cramer classes, SML values, and
IUPAC names) and writes them into six pre-defined tables inside the Word template.

Table layout
------------
+----------+-------------------------------------------------------+
| Template | Content                                               |
| position |                                                       |
+==========+=======================================================+
| Table 1  | FCM-listed substances grouped by CAS or FCM group.   |
|          | Columns: RT, name, CAS, FCM-name, classification,    |
|          | quantity (with uncertainty factor), SML, pass/fail.  |
+----------+-------------------------------------------------------+
| Table 2  | Toxicological classification summary for FCM entries.|
+----------+-------------------------------------------------------+
| Table 3  | Non-FCM / ungrouped substances with CAS.              |
|          | Optional CAS column controlled by ``cas_column`` flag.|
+----------+-------------------------------------------------------+
| Table 4  | Toxicological classification for ungrouped entries.  |
+----------+-------------------------------------------------------+
| Table A3 | Alkane substances detected in the sample.            |
+----------+-------------------------------------------------------+
| Table A4 | Unidentified substances (no CAS and no name).        |
+----------+-------------------------------------------------------+

Key public functions
--------------------
export_document(xlsx_path, out_path, text_changed, template, **kwargs)
    Orchestrates the entire export; called once per reporting run.

prepare_table1_data / prepare_table3_data
    Pre-process and group substance lists before handing data to the renderers.

render_table1 / render_table2 / render_table3 / render_table4 /
render_table_a3 / render_table_a4
    Each function populates one Word table via ``python-docx`` cell-level API.

_set_fcm_fields(s, df_eu, df_eu_grp, language, mapping_fcm_group, template_type)
    Enriches a substance object with FCM group name, SML value, and
    classification text from the EU2011 regulation look-up tables.

Uncertainty factor application
-------------------------------
Before writing quantities to tables, the analytical concentration (mg/L) is
multiplied by ``uf_real`` (typically ``1/6``) to obtain a realistic migration
estimate and by ``uf_exposition`` for exposure-adjusted values.

Output file naming
------------------
``<input_xlsx_basename>_<template_type>_<language>_<YYYYMMDD_HHMMSS>_results.docx``
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import time
from datetime import datetime
from chemical_classification import is_fcm_group, is_alkan, is_group
from rounding import list_to_str
from utilityFuncs.util import (
    set_table_cell_text,
    set_table_cell_text_fts,
    get_iupac_name,
    load_cramer_TTC_mapping,
    load_fcm_group_mapping
)

def export_document(xlsx_path, out_path, text_changed, template, **kwargs):
    """
    Generates and exports a Word document report based on analytical substance data.

    This function loads a DOCX template, populates multiple predefined tables 
    (Tables 1-4 and A3-A4) with processed data, saves intermediary table exports 
    to Excel for debugging, and then writes the completed report to disk.
    
    Notes:
    - Assumes set_table_cell_text, render_table*, and prepare_table* functions are defined elsewhere and handle table structure assumptions.
    - Relies on a specific Word template with tables at fixed positions (e.g. doc.tables[0] for Table 1).
    - Can be integrated into GUI applications due to the text_changed.emit(...) usage.
    
    Args:
        xlsx_path (str): Path to the input Excel file, used for naming the output file.
        out_path (str): Directory to write the final Word report.
        text_changed (signal): Qt signal or callback for updating status messages in a GUI.
        template (str): Base name of the DOCX template (without `.docx` extension).
        **kwargs: Dictionary of keyword arguments containing required data and configuration:
            - substances (list): List of substance objects.
            - isd (set): Set of CAS numbers to exclude (identified substances).
            - language (str): 'DE' or 'EN' for report language.
            - mapping_fcm_group (dict): FCM group-to-metadata mapping.
            - passed_yes (str): Status label for substances that pass evaluation.
            - passed_maybe (str): Label for uncertain/passable substances.
            - passed_no (str): Label for substances that fail.
            - uf_real (float): Realistic uncertainty factor.
            - uf_exposition (float): Exposure uncertainty factor.
            - cur_ft_dict (dict): Footnote tracking dictionary by table.
            - cas_column (bool): Whether Table 3 should include a CAS column.
            - cramer_mapping_group (dict): Group-based Cramer classifications.
            - cramer_mapping (dict): Substance-based Cramer classifications.
            - alkan (list): List to collect alkane-type substances.
            - unidentified (list): List of unidentified substances.

    Side Effects:
        - Writes intermediate Excel files to the `DEBUG` folder.
        - Emits status messages via `text_changed`.
        - Saves the final Word document to `out_path`.

    """
    doc = Document(f'{template}')
    if kwargs['language'] == "EN":
        kwargs['cramer_TTC_mapping'] = load_cramer_TTC_mapping(kwargs['db_path'], decimal_sep=".")
    else:
        kwargs['cramer_TTC_mapping'] = load_cramer_TTC_mapping(kwargs['db_path'])
    kwargs['mapping_fcm_group'] = load_fcm_group_mapping(kwargs['db_path'])
    # --- Table 1: FCM substances ---
    t1 = prepare_table1_data(
        kwargs['substances'], kwargs['isd'], kwargs['language'],
         kwargs['passed_yes'], kwargs['passed_maybe'],
        kwargs['uf_real'], kwargs['cur_ft_dict']['Table 1'],
        kwargs['df_eu'], kwargs['df_eu_grp'], kwargs['mapping_fcm_group'],
        kwargs['template_type']
    )
    render_table1(doc.tables[0], t1, text_changed)


    # --- Table 2: Toxicological classification of FCMs ---
    render_table2(doc.tables[1], kwargs['language'], t1['entries'], text_changed)

    # --- Table 3: Non-FCMs or ungrouped CAS substances ---
    t3 = prepare_table3_data(
        kwargs['substances'], kwargs['isd'], kwargs['alkan'], kwargs['cas_column'],
        kwargs['language'], kwargs['uf_real'], kwargs['uf_exposition'],
        kwargs['passed_maybe'], kwargs['passed_no'], kwargs['passed_yes'],
        kwargs['cramer_TTC_mapping'], kwargs['cur_ft_dict']['Table 3'],
        kwargs['K_gewicht']
    )
    render_table3(doc.tables[2], t3, text_changed)

    # --- Table 4: Toxicological classification of ungrouped substances ---
    render_table4(doc.tables[3], kwargs['language'], t3['entries'], text_changed)

    # --- Table A3: Alkanes ---
    render_table_a3(
        doc.tables[6], kwargs['alkan'], kwargs['language'],
        kwargs['uf_real'], kwargs['uf_exposition'], kwargs['passed_yes'], kwargs['passed_no'],
        kwargs['cramer_TTC_mapping'],kwargs['K_gewicht'],
        text_changed
    )

    # --- Table A4: Unidentified substances ---
    render_table_a4(
        doc.tables[7], kwargs['unidentified'], kwargs['language'],
        kwargs['uf_real'], kwargs['uf_exposition'],
        kwargs['cramer_TTC_mapping'],kwargs['K_gewicht'],
        text_changed
    )

    # --- Save final report ---
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    outname = xlsx_path.replace(".xlsx", f"_{kwargs['template_type']}_{kwargs['language']}_{timestamp}_results.docx")
    msg = f'Writing {out_path}\\{os.path.basename(outname)}'
    print(msg); text_changed.emit(msg)
    doc.save(f'{out_path}\\{os.path.basename(outname)}')
    print("Finished.... Exiting now")
    

def _set_fcm_fields(s, df_eu, df_eu_grp, language, mapping_fcm_group, template_type):
    """
    Populates FCM-related metadata fields on a substance object using the provided mapping.

    This function sets the FCM group name and toxicological classification (CL)
    in the appropriate language (German or English), based on the given FCM group.

    Summary:
    - This helper ensures that each substance is enriched with localized classification and name data from the mapping.
    - It is typically used during FCM group assignment in preparation for rendering Table 1.
    
    Args:
        s: A substance object with mutable attributes.
        group (str): The name or ID of the FCM group to which the substance belongs.
        mapping (dict): Dictionary mapping FCM group names to their associated metadata.
                        Expected keys per group: 'Name_DE', 'Name_EN', 'CL_DE', 'CL_EN'.
        language (str): 'DE' for German or 'EN' for English; determines which fields are used.

    Modifies:
        The substance object `s` in-place by assigning:
            - s.fcm_name
            - s.cl_de or s.cl_en (depending on language)
    """
    
    match = df_eu.loc[df_eu["FCM"] == s.fcm]
    col_name = "Name_EN" if language == "EN" else "Name_DE"
    # if language == 'EN':
    
    name_fcm = match[col_name].iloc[0] if not pd.isna(match[col_name].iloc[0]) else None
    grp_nr_from_eu_sheet = match["SML_Group"].iloc[0] if not pd.isna(match["SML_Group"].iloc[0]) else None
    
    if not match.empty and name_fcm:# if match found in EU2011 and name exists
        s.fcm_name = name_fcm # FCM name from EU2011
        
        sml_value = match["SML"].iloc[0] if not pd.isna(match["SML"].iloc[0]) else None
        if sml_value:
            s.sml = sml_value
        
        elif not sml_value and grp_nr_from_eu_sheet:# if no SML exists but group nr exists
            if grp_nr_from_eu_sheet.count("(") == 1:
                grp_nr = grp_nr_from_eu_sheet.replace("(", "").replace(")", "") if grp_nr_from_eu_sheet else None
                if grp_nr:
                    filtered = df_eu_grp.loc[df_eu_grp["Group_nr"] == int(grp_nr), "SML"]
                    if not filtered.empty:
                        s.sml = filtered.values[0]  # or use .item()
                    else:
                        s.sml = "-"  # if no grp nr exists
            elif grp_nr_from_eu_sheet.count("(") > 1:
                grp_nrs = [int(g.replace("(", "").replace(")", "")) for g in grp_nr_from_eu_sheet.split("\n")] # split multiple group numbers using ")("
                if grp_nrs:
                    filtered = df_eu_grp.loc[df_eu_grp["Group_nr"].isin(grp_nrs), "SML"]        
                    if not filtered.empty:
                        s.sml = min(filtered.values) # use the lowest SML if multiple group numbers exist
                    else:
                        s.sml = "-"  # if no grp nr exists

                    
    # elif language == 'DE':
    #     if not match.empty and not pd.isna(match["Name_DE"].iloc[0]):
    #         s.fcm_name = match["Name_DE"].iloc[0]
    #         if not pd.isna(match["SML"].iloc[0]) :
    #             s.sml = match["SML"].iloc[0]
    #         elif pd.isna(match["SML"].iloc[0]) and not pd.isna(match["SML_Group"].iloc[0]):
    #             grp_nr_from_eu_sheet = match["SML_Group"]
    #             grp_nr = grp_nr_from_eu_sheet.iloc[0].replace("(", "").replace(")", "") if not grp_nr_from_eu_sheet.empty else None
    #             if grp_nr:
    #                 filtered = df_eu_grp.loc[df_eu_grp["Group_nr"] == int(grp_nr), "SML"]
    #             if not filtered.empty:
    #                 s.sml = filtered.values[0]  # or use .item()
    #             else:
    #                 s.sml = "-"  # if no grp nr exists    
    else:
        s.fcm_name = "EUFCMNamenotfound"
        print(f"Skipping {s.cas}")
    
    if s.fcm:
        try:
            fcm_key = str(s.fcm)
            info = mapping_fcm_group[fcm_key]
            if language == 'DE' and ((template_type == mapping_fcm_group[fcm_key]['Type_Flag']) | (mapping_fcm_group[fcm_key]['Type_Flag'] == "A")):
                s.fcm_name = info['Name_DE']
                s.cl_de = info['CL_DE']
                s.cas = "-"  # FCMs do not have CAS numbers

            elif language == "EN" and ((template_type == mapping_fcm_group[fcm_key]['Type_Flag']) | (mapping_fcm_group[fcm_key]['Type_Flag'] == "A")):
                    s.fcm_name = info['Name_EN']
                    s.cl_en = info['CL_EN']
                    s.cas = "-"  # FCMs do not have CAS numbers
                    
                    
        except KeyError:
            print(f"FCM group {s.fcm} not found in mapping.")
            match = df_eu.loc[df_eu["FCM"] == fcm_key]
            if language == 'EN':
                if not match.empty and not pd.isna(match["Name_EN"].iloc[0]):
                    s.fcm_name = match["Name"].iloc[0]
                    s.cl_en = "No classification in EU2011 found"
            elif language == 'DE':
                if not match.empty and not pd.isna(match["Name_DE"].iloc[0]):
                    s.fcm_name = match["Name_DE"].iloc[0]
                    s.cl_de =  "Keine Klassifizierung in EU2011 gefunden"
            else:
                s.fcm_name = "FCM name not found in EU"
                print(f"Skipping {s.cas}")
        
# --- Table 1 Helpers ---
def prepare_table1_data(substances, isd, language,
                        passed_yes, passed_maybe, uf_real, start_ft, 
                        df_eu, df_eu_grp, mapping_fcm_group,
                        template_type):
    """
    Prepares and organizes data for Table 1, grouping substances by CAS or FCM group
    and setting necessary metadata for rendering.

    This function:
        - Filters out substances in the inclusion substance dataset (isd).
        - Groups substances with FCM designations by CAS number.
        - Groups substances without explicit FCM designation but matching known FCM groups
          via SMILES-based lookup.
        - Populates key classification fields for matched FCM groups using a mapping dictionary.
        - Tracks the first entry per group for table row creation.
    
    Notes:
        - FCM group resolution is attempted only if s.fcm is not already set.
        -  _set_fcm_fields fills localized names/classifications from the group mapping.
        - entries acts as the ordered list for rendering table rows — one per group or CAS.

    Args:
        substances (list): List of substance objects, each expected to have at least:
            - .cas (str)
            - .fcm (str or None)
            - .canonical_smiles (str or None)
        isd (set): Set of CAS numbers (strings) to exclude (e.g. included in other tables).
        language (str): Either 'DE' or 'EN'; used for localized field mapping.
        mapping_fcm_group (dict): Dictionary mapping FCM group names to structured metadata.
        passed_yes (str): Label indicating a substance has passed toxicological assessment.
        passed_maybe (str): Label for substances with potential but uncertain pass status.
        uf_real (float): Realistic uncertainty factor used in later exposure calculation.
        start_ft (str): Starting character (e.g. 'a') for footnote labeling.

    Returns:
        dict: Context dictionary containing:
            - 'data': A dict mapping CAS or FCM group keys to lists of substance objects.
            - 'entries': A list of first representative substance per group (for table rows).
            - 'language', 'passed_yes', 'passed_maybe', 'uf_real', 'start_ft': Passed-through values.
    """
    table_data = {} # Mapping of CAS or FCM group keys to lists of substances
    entries = [] # First entries per group for row generation
    for s in substances:
        # Skip substances in the inclusion dataset (isd)
        if s.cas.lstrip('0') in isd:
            continue
        
        # If FCM is already defined, set group == s.fcm
        grp=None
        if s.fcm:
            grp = s.fcm
            
        # Otherwise, try use SMILES to set FCM group
        elif s.canonical_smiles:
            
            grp = is_fcm_group(s.canonical_smiles)
        
        # if a grp is set
        if grp:
            
            # Assign FCM group to the substance
            if not s.fcm:
                s.fcm = grp
            table_data.setdefault(grp, []).append(s)
            
            # Populate classification fields from mapping, localized by language
            _set_fcm_fields(s, df_eu, df_eu_grp, language, mapping_fcm_group, template_type)
            
            # Only append the first entry for rendering
            if len(table_data[grp]) == 1:
                entries.append(s)
    return {
        'data': table_data,
        'entries': entries,
        'language': language,
        'passed_yes': passed_yes,
        'passed_maybe': passed_maybe,
        'uf_real': uf_real,
        'start_ft': start_ft
    }


def render_table1(table, ctx, text_changed):
    """
    Populates Table 1 in a Word document with grouped FCM substance data.

    Each row in the table represents a group of substances that share an FCM assignment
    or a common CAS number. The table includes information such as CAS, FCM group,
    retention time (RT), qualification level, total quantity, uncertainty-adjusted quantity,
    specific migration limits (SML), and evaluation status. Footnotes are appended
    for substances where applicable.
    
    Additional Notes:
    Helper functions like list_to_str, _assign_footnotes, set_table_cell_text_fts, and _insert_footnote_block are assumed to handle formatting 
    and Word table/text operations.
    
    Args:
        table (docx.table.Table): The Word table object to be populated.
        ctx (dict): Context dictionary containing all required data and labels:
            - 'data' (dict): Maps either an FCM group index (int) or a CAS number (str)
              to a list of substance-like objects.
            - 'passed_yes' (str): Label used for substances passing the evaluation.
            - 'passed_maybe' (str): Label used for substances with uncertain classification.
            - 'uf_real' (float): Realistic uncertainty factor applied to quantity values.
            - 'start_ft' (int): Starting index for footnote tracking.
            - 'language' (str): Report language ('DE' or 'EN').

    Raises:
        ValueError: If the Word table does not have the expected number of columns.

    Side Effects:
        - Adds rows to the provided Word table.
        - Inserts a block of footnotes below the table via `_insert_footnote_block`.
    """    
    # a dictionary mapping either an FCM group index (int) or a CAS number (str) to a list of substance-like objects.
    data = ctx['data']
    
    # localized labels or booleans indicating pass/fail status.
    p_yes = ctx['passed_yes']
    p_maybe = ctx['passed_maybe']
    
    # a numeric uncertainty factor used to adjust quant values.
    uf = ctx['uf_real']
    
    # for tracking footnotes.
    cur_ft = ctx['start_ft']
    fts_map = {}
    
    
    lang = ctx['language']
    # start writing at the first data row - skip header row
    row = 1

    for key, group in data.items():
        # make sure this row exists; Word tables start with only header (row 0)
        if row >= len(table.rows):
            table.add_row()
        
        if len(group) > 1:# if FCM with multiple CAS - FCM Group
            cas = '-'
            print(f"{group[0].sml} replaced by '-' for group {group[0]} containing > 1 element")
            sml = "-"
        else: # FCM is with single CAS
            cas = group[0].cas.lstrip('0')
            try:
                sml = str(s0.sml)
            except UnboundLocalError:
                print(f"Something wrong with the SML for {group[0]} and FCM {group[0].fcm}")
                sml = f"SML NOT FOUND for FCM {group[0].fcm}"
        
        #For FCM  Use the first substance as a representative, if len(grp)>1
        s0 = group[0]
        fcm_name = s0.fcm_name
        if isinstance(s0.sml, (int, float)):
            sml = str(s0.sml)
        else:
            print(f"{s0.sml} replaced by '-' for {s0}")
            sml = "-"
        # Format retention times (RT) and qualification level
        rt = list_to_str([s.rt for s in group], 'nachkomma') 
        
        # Minimum qualification level, prefixed with "≥" if multiple substances.
        qual = f"≥ {min(int(s.qual) for s in group)}" if len(group) > 1 else f"{int(s0.qual)}"
        # Quantitative values (raw and adjusted)
        total_quant = sum(s.quant for s in group)
        quant =  list_to_str([total_quant], 'significant', keep_signf=2)
        quant_real = list_to_str([total_quant * uf ], 'significant', keep_signf=2)
        
        # Determine pass/fail status (based on SML if available)
        passed = p_yes if sml == '-' or sum(s.quant for s in group) < float(s0.sml or 0) else p_maybe
        
        # Assign and track footnotes
        row_fts, cur_ft, fts_map = _assign_footnotes(s0, lang, fts_map, cur_ft)
        
        # Populate the row with formatted values
        set_table_cell_text_fts(table, row, 0, fcm_name, row_fts)
        # if len(rt)>1:
        #     try:
        #         cas = s0.cas.lstrip('0')
        #     except:
        #         print("Something wrong with the CAS during rt replacement")
        set_table_cell_text(table, row, 1, cas, center=True)
        set_table_cell_text(table, row, 2, str(s0.fcm))
        set_table_cell_text(table, row, 3, rt, center=True)
        set_table_cell_text(table, row, 4, qual, center=True)
        set_table_cell_text(table, row, 5, quant, center=True)
        set_table_cell_text(table, row, 6, quant_real, center=True)
        set_table_cell_text(table, row, 7, sml, center=True)
        set_table_cell_text(table, row, 8, passed, center=True)
        row += 1

    # Append footnote block below the table
    _insert_footnote_block(table._parent, fts_map, 'FT1')
    text_changed.emit("Table 1 finished...")
    print("table 1 finished...")

def _assign_footnotes(s, language, existing_map, start_ft):
    """
    Assigns footnotes to a substance based on its footnote text in the specified language.
    
    Footnotes are tracked using a letter key (e.g., 'a', 'b', 'c'), starting from `start_ft`.
    If a footnote text has already been assigned a key in `existing_map`, it reuses the key.
    Otherwise, it assigns the next available letter and updates the map.

    Args:
        s: A substance-like object with attributes `ft_en` and `ft_de`, each containing
           either a string or a list of footnote texts in English or German.
        language (str): Language code, either 'EN' or 'DE', to select the appropriate footnotes.
        existing_map (dict): A mapping of previously assigned footnote keys to their texts.
        start_ft (str): The current footnote label (e.g., 'a', 'b', etc.) to start assigning from.

    Returns:
        tuple: A tuple containing:
            - row_map (dict): A mapping of footnote keys to footnotes for the current substance.
            - cur (str): The next available footnote label after the current assignment.
            - existing_map (dict): The updated global map of footnote labels to texts.
    """
    fts = s.ft_de if language == 'DE' else s.ft_en
    fts = [fts] if isinstance(fts, str) else fts
    row_map = {}; cur = start_ft
    for ft in fts:
        if ft in existing_map.values():
            key = list(existing_map.keys())[list(existing_map.values()).index(ft)]
            row_map[key] = ft
        else:
            row_map[cur] = ft; existing_map[cur] = ft
            cur = chr(ord(cur) + 1)
    return row_map, cur, existing_map


def _insert_footnote_block(doc_parent, fts_map, marker):
    """
    Inserts a block of footnotes into a Word document at the location of a given marker paragraph.

    Uses a pre-defined Word style (e.g., 'Footnote') with a hanging indent and tab stop set at 0.4″ (567 twips),
    so that each footnote aligns correctly via '\t'.
    """
    # --- Find and clear the marker paragraph ---
    for p in doc_parent.paragraphs:
        if p.text == marker:
            p.text = ''
            break

    # --- Insert footnotes before the cleared marker paragraph ---
    for ft, txt in fts_map.items():
        np = p.insert_paragraph_before()
        np.style = 'Footnote'  # Assumes Word style with desired tab stop and indent
        np.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Superscript label
        r1 = np.add_run()
        r1.text = ft
        r1.font.name = 'Arial'
        r1.font.size = Pt(5)
        r1.font.superscript = True

        # Tab + footnote text
        r2 = np.add_run()
        r2.text = f"\t{txt}"
        r2.font.name = 'Arial'
        r2.font.size = Pt(5)


# --- Table 2 ---
def render_table2(table, language, entries, text_changed):
    """
    Populates Table 2 in a Word document with substance classification data.

    This function iterates over a list of substance-like `entries`, and for each substance:
      - Retrieves the FCM name and CAS number.
      - Selects classification text based on the provided language ('EN' or 'DE').
      - Inserts a new row into the given Word table and populates its columns accordingly.

    Table columns are expected to be:
        [0] FCM Name
        [1] CAS Number (centered)
        [2] Classification (single or multiple entries)
    
    Notes:
        - Assumes the use of a helper function set_table_cell_text(table, row, col, text, center=False), 
        likely a wrapper to simplify setting Word table cell contents and formatting.
        - Assumes the first row (row 0) is reserved for column headers.
        - Handles both string and list classification formats.
        - Ensures CAS numbers are stripped of leading zeros for formatting consistency.
        
    Args:
        table: A python-docx table object into which the data will be written.
        language (str): Language code ('EN' or 'DE') to choose classification labels.
        entries (list): A list of substance-like objects, each with:
                        - fcm_name
                        - cas
                        - cl_en / cl_de (classification list or string)

    Returns:
        None. The function modifies the table in-place.
    """
    row = 1 # Start inserting from the second row (assuming row 0 is header)
    for s in entries:
        # Select classification(s) in the chosen language
        cls = s.cl_de if language == 'DE' else s.cl_en
        cls = [cls] if isinstance(cls, str) else cls # Normalize to list
        
        # Handle missing FCM name
        name = s.fcm_name or '[no FCM name]'
        
        # Add a new row to the table
        table.add_row()
        
        # Populate the row cells
        set_table_cell_text(table, row, 0, name)
        set_table_cell_text(table, row, 1, s.cas.lstrip('0'), center=True)
        set_table_cell_text(table, row, 2, cls)
        
        row += 1
    print("Table 2 finished...")
    text_changed.emit("Table 2 finished...")

# --- Table 3 ---
def prepare_table3_data(substances, isd, alkan, cas_column, 
                        language, uf_real, uf_expo, 
                        p_maybe, p_no, p_yes,
                        cramer_TTC_mapping, start_ft, K_gewicht):
    """
    Prepares and filters data for populating Table 3 in a regulatory Word report.

    This function filters a list of `substances` into grouped and ungrouped entries
    based on FCM relevance, chemical structure (SMILES), and grouping logic.
    It excludes irrelevant or handled substances (like ISDs and alkanes), identifies
    structurally similar groups, and collects ungrouped substances for individual listing.
    
    Summary of Logic:
        Skips:
            - CAS numbers listed in isd
            - Substances already associated with an FCM or FCM group

        Collects:
            - Alkanes separately via is_alkan
            - Grouped substances using is_group()
            - Ungrouped ones by CAS (included in the final entries list)
            
    Args:
        substances (list): List of substance-like objects, each with at least:
                           - cas
                           - fcm
                           - canonical_smiles
        isd (set): Set of CAS numbers (as strings, no leading zeros) to exclude.
        alkan (list): A list to collect substances identified as alkanes.
        cas_column (bool): Whether to display CAS numbers in the final table.
        language (str): 'EN' or 'DE', used for group name resolution.
        uf_real (float): Realistic uncertainty factor for quantitative adjustments.
        uf_expo (float): Exposure-related uncertainty factor (not used here, but passed through).
        p_maybe, p_no, p_yes (str): Text labels for regulatory status decisions.
        cramer_group_map (dict): Mapping for Cramer classification footnotes per group.
        cramer_map (dict): Mapping for individual Cramer classifications.
        start_ft (str): Starting label for footnote assignment (e.g., 'a').

    Returns:
        dict: A context dictionary with the following keys:
            - 'data': dict mapping group names or CAS numbers to lists of substances
            - 'entries': list of ungrouped, non-alkan, non-ISD substances
            - 'cas_column': original flag
            - 'language': selected language
            - 'uf_real': uncertainty factor
            - 'uf_expo': exposure uncertainty factor
            - 'p_yes', 'p_maybe', 'p_no': status labels
            - 'cramer_group_map': mapping for group-level footnotes
            - 'cramer_map': mapping for individual substance-level Cramer footnotes
            - 'start_ft': updated footnote label state
    """

    data = {} # Will hold group or CAS → list of substances
    entries = []  # Will collect ungrouped, CAS-identified substances
    for s in substances:
        # Skip substances already handled as ISDs
        if s.cas.lstrip('0') in isd: continue
        
        # Skip substances with an assigned FCM or identified FCM group
        if s.fcm or (s.canonical_smiles and is_fcm_group(s.canonical_smiles)): continue
        
        # Identify and collect alkanes separately
        if is_alkan(s.canonical_smiles): 
            alkan.append(s)
            continue
        
        # Try to assign the substance to a structural group
        grp = is_group(s.canonical_smiles, language)
        if grp: 
            data.setdefault(grp, []).append(s)
        else:
            # No group found: use CAS as key and track in entries list
            key = s.cas
            data.setdefault(key, []).append(s)
            entries.append(s)
    return {
        'data': data, 'entries': entries, 'cas_column': cas_column,
        'language': language, 'uf_real': uf_real, 'uf_expo': uf_expo,
        'p_yes': p_yes, 'p_maybe': p_maybe, 'p_no': p_no,
        'cramer_TTC_mapping': cramer_TTC_mapping,
        'start_ft': start_ft, 'K_gewicht': K_gewicht
    }
    
    
def ensure_table_has_rows(table, required_rows):
    """
    Ensures the table has at least the required number of rows.

    Args:
        table: A `python-docx` table object.
        required_rows: The minimum number of rows the table should have.

    Returns:
        None. Adds rows to the table if necessary.
    """
    current_rows = len(table.rows)
    for _ in range(required_rows - current_rows):
        table.add_row()

def render_table3(table, ctx, text_changed):
    """
    Populates Table 3 in a Word document with detailed information on substances not in FCM groups,
    including TTC evaluation, Cramer classification, and exposure estimates.

    This function uses a context dictionary (`ctx`) to extract data, configuration, and mappings
    required for rendering each row of the table, and adds footnotes as needed.
    
    Summary:
    - Each row includes fields like name, CAS, RT, quantity, Cramer class, TTC, exposure estimate, and pass/fail classification.
    - Cramer and TTC data are retrieved from mappings, either by group or individual classification.
    - A footnote block is inserted after all rows via _insert_footnote_block.
    
    Args:
        table: A `python-docx` table object to be populated.
        ctx (dict): Context dictionary with the following keys:
            - 'data': dict mapping group names or CAS numbers to lists of substances
            - 'language': 'EN' or 'DE'
            - 'uf_real': uncertainty factor for real exposure
            - 'uf_expo': multiplier for exposure adjustment
            - 'p_yes', 'p_maybe', 'p_no': status labels
            - 'cramer_group_map': dict mapping group keys to Cramer classification info
            - 'cramer_map': dict mapping Cramer codes to classification info
            - 'start_ft': initial footnote label (e.g. 'a')
            - 'cas_column': bool, whether to include the CAS column.

    Raises:
        ValueError: If the table template has an incorrect number of columns.

    Returns:
        None. The function modifies the table in-place.
    """
    data = ctx['data']
    lang = ctx['language']
    uf = ctx['uf_real']
    expo = ctx['uf_expo']
    gewicht = ctx["K_gewicht"]
    p_yes = ctx['p_yes']
    p_maybe = ctx['p_maybe']
    p_no = ctx['p_no']
    cramer_TTC_mapping = ctx['cramer_TTC_mapping']
    cur_ft = ctx['start_ft']
    fts_map = {}
    row = 1

    # Ensure the table has enough rows for the data
    ensure_table_has_rows(table, len(data) + 1)  # +1 for the header row

    for grp_name, group in data.items():
        
        s0 = group[0]  # Representative substance for the group
        # --- Determine CAS and Cramer classification ---

        if s0.cramer:
            cramer = s0.cramer
            ttc = cramer_TTC_mapping[s0.cramer]['TTC']
            ttc_adj = round(ttc * gewicht, 2)
        else:
            cas = s0.cas.lstrip('0')
            cramer = '?'
            ttc = None
            ttc_adj = ttc

            
        # --- Compute and format values ---
        rt = list_to_str([s.rt for s in group], 'nachkomma')
        qual = f"≥ {min(int(s.qual) for s in group)}" if len(group) > 1 else f"{int(s0.qual)}"
        quant = list_to_str([s.quant for s in group], 'significant', keep_signf=2)
        quant_real = list_to_str([s.quant * uf for s in group], 'significant', keep_signf=2)
        exposure_vals = [s.quant * uf * expo for s in group]
        exposure = list_to_str(exposure_vals, 'significant', keep_signf=2)
        passed = p_maybe if ttc_adj is None else (p_no if any(e >= ttc_adj for e in exposure_vals) else p_yes)

        # override cas if "passed"
        if cramer != "?" and ttc_adj is not None:
            cas = s0.cas.lstrip('0')

        # --- Assign footnotes ---
        row_fts, cur_ft, fts_map = _assign_footnotes(s0, lang, fts_map, cur_ft)
        
        # group name and CAS handling
        if len(group) > 1 and grp_name.isalpha():
            name = grp_name
            cas = None  # No CAS for groups
            print(name, cas)
            time.sleep(0.1)  # Small delay to avoid
        else:
            name = get_iupac_name(lang, s0)
            cas = s0.cas.lstrip('0')
        
        # If CAS column is not used, just set the name, else add CAS
        if cas:
            name = name + " " +"\n" + "(" + cas + ")"
        else:
            name = name + " " +"\n"
            
        # --- Fill in the row ---
        col = 0
        if ctx['cas_column']:
            name = name.replace('" " +"\n" + "(" + cas + ")', '" " +"\n"')
            set_table_cell_text_fts(table, row, col, name, row_fts)
            col += 1
            if not cas:
                cas = "-"
            set_table_cell_text(table, row, col, cas, center=True)
        else:
            set_table_cell_text_fts(table, row, col, name, row_fts)
            col += 1  # Skip over CAS column if not used
        col += 1; set_table_cell_text(table, row, col, rt, center=True)
        col += 1; set_table_cell_text(table, row, col, qual, center=True)
        col += 1; set_table_cell_text(table, row, col, quant, center=True)
        col += 1; set_table_cell_text(table, row, col, quant_real, center=True)
        col += 1; set_table_cell_text(table, row, col, exposure, center=True)
        col += 1; set_table_cell_text(table, row, col, cramer, center=True)
        ttc_adj = list_to_str([ttc_adj], 'significant') if ttc_adj is not None else None
        col += 1; set_table_cell_text(table, row, col, ttc_adj or '?', center=True)
        col += 1; set_table_cell_text(table, row, col, passed, center=True)

        row += 1

    # Add footnote block below the table
    _insert_footnote_block(table._parent, fts_map, 'FT3')
    text_changed.emit("Table 3 finished...")
    print("Table 3 finished...")

# --- Table 4 ---
def render_table4(table, language, entries, text_changed):
    """
    Populates Table 4 in a Word document with data on substances that have 
    toxicological classifications but are not included in FCM groups.

    Each row contains:
        - The substance name (IUPAC or fallback name, language-dependent)
        - The CAS number (with leading zeroes removed)
        - The toxicological classification (language-dependent)
    
    Summary:
    - Table 4 includes non-FCM substances with toxicological concern.
    - Classification and names are selected based on the active language.
    - Each substance appears on its own row with consistent formatting.
    
    Args:
        table: A `python-docx` table object to populate.
        language (str): 'DE' for German or 'EN' for English text.
        entries (list): List of substance-like objects, each expected to have:
            - .cl_de / .cl_en: Toxicological classification(s)
            - .iupac_name_de / .iupac_name_en: IUPAC names
            - .name: General name fallback
            - .cas: CAS number string

    Returns:
        None. The function modifies the `table` in-place.
    """
    keys_done = set()    
    row = 1
    for s in entries:
        # Select classification based on language
        cls = s.cl_de if language=='DE' else s.cl_en
        cls = [cls] if isinstance(cls, str) else cls # Normalize to list if needed
        
        # Use IUPAC name if available, otherwise fallback to general name
        iupac = s.iupac_name_de or s.name if language=='DE' else s.iupac_name_en or s.name
        
        cas_nr = s.cas.lstrip('0')
        if cas_nr:
            if s.cas.lstrip('0') in keys_done:
                continue
            else:
                keys_done.add(s.cas.lstrip('0'))
        else:
            if iupac in keys_done:
                continue
            else:
                keys_done.add(iupac)
            
        # Add and fill row
        table.add_row()
        set_table_cell_text(table, row, 0, iupac)
        set_table_cell_text(table, row, 1, s.cas.lstrip('0'))
        set_table_cell_text(table, row, 2, cls)
        row += 1
    
    text_changed.emit("Table 4 finished...")
    print("Table 4 finished...")
    
# --- Table A3 ---
def render_table_a3(table, alkanes, language, uf_real, uf_expo, passed_yes, passed_no,
                    cramer_TTC_mapping, gewicht, text_changed):
    """
    Renders Table A3 in a Word document, showing exposure calculations and 
    compliance evaluation for identified alkanes.

    Each row includes:
        - A fixed label 'Alkane'
        - Retention time (language-specific decimal format)
        - Reported quantity
        - Quantity after applying the real uncertainty factor (UF_real)
        - Quantity after applying both UF_real and UF_expo
        - Fixed TTC threshold value
        - Pass/fail result depending on whether calculated exposure is below TTC

    Args:
        table: A `python-docx` table object to populate.
        alkanes (list): List of alkane-like objects, each expected to have:
            - .rt: Retention time (float)
            - .quant: Quantity (float)
        language (str): Either 'DE' (German) or 'EN' (English), affects decimal formatting.
        uf_real (float): Realistic uncertainty factor to adjust the quantity.
        uf_expo (float): Exposure factor to further adjust the quantity.
        passed_yes (str): Label to use when substance passes TTC check.
        passed_no (str): Label to use when substance fails TTC check.

    Returns:
        None. Modifies the table in-place.
    """
    #Retention

    # Cramer classification and TTC mapping
    # Assuming ctx is a dictionary with 'cramer_map' and 'cramer_mapping_group
    
    row = 1
    # ttc_val = 1.8 # TTC threshold in µg (Fixed)
    ttc_key = 'I'  # Default key, replace with dynamic logic if needed
    if ttc_key not in cramer_TTC_mapping:
        raise KeyError(f"Key '{ttc_key}' not found in cramer_TTC_mapping")

    ttc = cramer_TTC_mapping[ttc_key]['TTC']
    
    ttc_adj = ttc * gewicht
    for s in alkanes:
        # Determine pass/fail based on exposure vs TTC
        exposure = s.quant*uf_real*uf_expo
        passed = passed_yes if exposure < ttc_adj else passed_no
        
        # Format retention time based on language
        rt_txt = f"{s.rt:.2f}".replace('.', ',') if language=='DE' else f"{s.rt:.2f}"
        
        # Add and populate a new row
        table.add_row()
        fixed_label = 'Alkan' if language == 'DE' else 'Alkane'
        # Use the fixed label for the first column
        # Set the table cells with appropriate values
        set_table_cell_text(table, row, 0, fixed_label)
        set_table_cell_text(table, row, 1, rt_txt, center=True)
        set_table_cell_text(table, row, 2, list_to_str([s.quant], 'significant', keep_signf=2), center=True)
        set_table_cell_text(table, row, 3, list_to_str([s.quant*uf_real], 'significant', keep_signf=2), center=True)
        set_table_cell_text(table, row, 4, list_to_str([s.quant*uf_real*uf_expo], 'significant', keep_signf=2), center=True)
        set_table_cell_text(table, row, 5, list_to_str([ttc_adj], 'significant'), center=True)
        set_table_cell_text(table, row, 6, passed, center=True)
        row += 1
    print("Table a3 finished...")
    text_changed.emit("Table A3 finished...")
    
# --- Table A4 ---
def render_table_a4(table, unidentified, language, uf_real, uf_expo,cramer_TTC_mapping, gewicht, 
                    text_changed):
    """
    Populates Table A4 with unidentified substances, showing estimated exposure 
    and indicating that the toxicological assessment is inconclusive.

    Each row includes:
        - A fixed label: 'Unidentifiziert' or 'Unidentified'
        - Retention time (with language-specific decimal format)
        - Reported quantity
        - Quantity adjusted with UF_real
        - Quantity adjusted with UF_real and UF_expo
        - Fixed TTC value (0.00015 µg/kg)
        - Status: always 'nicht eindeutig' or 'inconclusive'
        
    Summary:
    - Table A4 deals with unidentified substances.
    - Exposure is calculated but always marked "inconclusive" due to lack of identity.
    - TTC value is conservatively fixed at 0.00015 µg/kg, following common regulatory practice.
    
    Args:
        table: A `python-docx` table object to populate.
        unidentified (list): List of unidentified substance-like objects, each expected to have:
            - .rt: Retention time (float)
            - .quant: Quantity (float)
        language (str): Either 'DE' (German) or 'EN' (English), affects label and decimal formatting.
        uf_real (float): Realistic uncertainty factor applied to quantity.
        uf_expo (float): Exposure factor applied after UF_real.

    Returns:
        None. The function modifies the table in-place.
    """

    row = 1
    # ttc_val = 0.00015 
    # TTC threshold for unidentified substances (µg/kg)
    ttc_key = 'GENOTOX'  # Default key, replace with dynamic logic if needed
    if ttc_key not in cramer_TTC_mapping:
        raise KeyError(f"Key '{ttc_key}' not found in cramer_TTC_mapping")

    ttc = cramer_TTC_mapping[ttc_key]['TTC'] 
    ttc_adj = ttc * gewicht
    
    for s in unidentified:
        # Fixed label and status depending on language
        text = 'Unidentifiziert' if language=='DE' else 'Unidentified'
        passed = 'nicht eindeutig' if language=='DE' else 'inconclusive'
        
        # Format retention time with comma in German
        rt_txt = f"{s.rt:.2f}".replace('.', ',') if language=='DE' else f"{s.rt:.2f}"
        # list_to_str([total_quant * uf ], 'significant')
        # Add a new row and populate each cell
        table.add_row()
        set_table_cell_text(table, row, 0, text)
        set_table_cell_text(table, row, 1, rt_txt, center=True)
        set_table_cell_text(table, row, 2, list_to_str([s.quant ], 'significant', keep_signf=2), center=True)
        set_table_cell_text(table, row, 3, list_to_str([s.quant*uf_real ], 'significant', keep_signf=2), center=True)
        set_table_cell_text(table, row, 4, list_to_str([s.quant*uf_real*uf_expo ], 'significant', keep_signf=2), center=True)
        set_table_cell_text(table, row, 5, list_to_str([ttc_adj], 'significant'), center=True)
        set_table_cell_text(table, row, 6, passed, center=True)
        row += 1
        
    print("Table a4 finished...")
    text_changed.emit("Table A4 finished...")